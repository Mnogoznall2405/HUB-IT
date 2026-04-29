from __future__ import annotations

import csv
import io
import json
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

from fastapi import UploadFile

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None

try:
    from openpyxl import Workbook  # type: ignore
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side  # type: ignore
    from openpyxl.utils import get_column_letter  # type: ignore
except Exception:  # pragma: no cover
    Alignment = None
    Border = None
    Font = None
    Workbook = None
    PatternFill = None
    Side = None
    get_column_letter = None

try:
    from reportlab.lib.pagesizes import A4  # type: ignore
    from reportlab.pdfbase import pdfmetrics  # type: ignore
    from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
    from reportlab.pdfgen import canvas  # type: ignore
except Exception:  # pragma: no cover
    A4 = None
    TTFont = None
    canvas = None
    pdfmetrics = None


SUPPORTED_GENERATED_FILE_FORMATS = {"csv", "xlsx", "docx", "pdf", "txt", "md", "json"}
MAX_GENERATED_FILES_PER_RUN = 10
MAX_GENERATED_FILE_BYTES = 10 * 1024 * 1024
MAX_GENERATED_FILE_ROWS_PER_SHEET = 50000

_CONTENT_TYPES = {
    "csv": "text/csv; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "json": "application/json; charset=utf-8",
    "md": "text/markdown; charset=utf-8",
    "pdf": "application/pdf",
    "txt": "text/plain; charset=utf-8",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

_PDF_FONT_NAME = "AIGeneratedUnicode"
_PDF_FONT_REGISTERED = False
_EXCEL_INVALID_SHEET_TITLE_RE = re.compile(r"[\x00-\x1f\\/*?:\[\]]+")
_EMPTY_EQUIPMENT_SERIAL = "\u2014"
_EQUIPMENT_TABLE_COLUMNS: tuple[tuple[str, str, tuple[str, ...]], ...] = (
    ("inv_no", "\u0418\u043d\u0432. \u043d\u043e\u043c\u0435\u0440", ("inv_no", "inventory_number", "inventory_no", "inv", "\u0438\u043d\u0432", "\u0438\u043d\u0432 \u043d\u043e\u043c\u0435\u0440", "\u0438\u043d\u0432\u0435\u043d\u0442\u0430\u0440\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440")),
    ("serial_no", "\u0421\u0435\u0440\u0438\u0439\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440", ("serial_no", "serial_number", "serial", "sn", "\u0441\u0435\u0440\u0438\u0439\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440", "\u0441\u0435\u0440\u0438\u0439\u043d\u0438\u043a")),
    ("type_name", "\u0422\u0438\u043f", ("type_name", "item_type", "type", "\u0442\u0438\u043f")),
    ("model_name", "\u041c\u043e\u0434\u0435\u043b\u044c", ("model_name", "model", "name", "\u043c\u043e\u0434\u0435\u043b\u044c")),
    ("owner_name", "\u0421\u043e\u0442\u0440\u0443\u0434\u043d\u0438\u043a", ("owner_name", "employee_name", "fio", "owner_fio", "user_fio", "\u0441\u043e\u0442\u0440\u0443\u0434\u043d\u0438\u043a", "\u0432\u043b\u0430\u0434\u0435\u043b\u0435\u0446")),
    ("status", "\u0421\u0442\u0430\u0442\u0443\u0441", ("status", "status_name", "item_status", "\u0441\u0442\u0430\u0442\u0443\u0441")),
    ("branch", "\u0424\u0438\u043b\u0438\u0430\u043b", ("branch", "branch_name", "filial", "company", "\u0444\u0438\u043b\u0438\u0430\u043b")),
    ("location", "\u041b\u043e\u043a\u0430\u0446\u0438\u044f", ("location", "location_name", "room", "cabinet", "\u043b\u043e\u043a\u0430\u0446\u0438\u044f", "\u043c\u0435\u0441\u0442\u043e", "\u043a\u0430\u0431\u0438\u043d\u0435\u0442")),
)
_EQUIPMENT_ALIAS_LOOKUP = {
    re.sub(r"[\s._-]+", " ", alias.casefold().replace("\u0451", "\u0435")).strip(" :"): canonical
    for canonical, _label, aliases in _EQUIPMENT_TABLE_COLUMNS
    for alias in aliases
}


class GeneratedFileError(ValueError):
    def __init__(
        self,
        message: str,
        *,
        error_code: str = "generated_file_error",
        field_path: str | None = None,
        suggested_fix: str | None = None,
    ) -> None:
        super().__init__(message)
        self.error_code = error_code
        self.field_path = field_path
        self.suggested_fix = suggested_fix

    def to_payload(self) -> dict[str, Any]:
        return {
            "error_code": self.error_code,
            "message": str(self),
            "field_path": self.field_path,
            "suggested_fix": self.suggested_fix,
        }


@dataclass(slots=True)
class GeneratedTableSpec:
    title: str = ""
    rows: list[list[str]] = field(default_factory=list)
    header_row_index: int = 1

    def to_dict(self) -> dict[str, Any]:
        return {
            "title": self.title,
            "rows": self.rows,
            "header_row_index": self.header_row_index,
        }


@dataclass(slots=True)
class GeneratedSheetSpec:
    title: str
    rows: list[list[str]] = field(default_factory=list)
    header_row_index: int = 1

    def to_dict(self) -> dict[str, Any]:
        return {
            "title": self.title,
            "rows": self.rows,
            "header_row_index": self.header_row_index,
        }


@dataclass(slots=True)
class GeneratedFileSpec:
    format: str
    file_name: str
    title: str = ""
    content: Any = None
    rows: list[list[str]] = field(default_factory=list)
    sheets: list[GeneratedSheetSpec] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    size_bytes: int = 0

    def to_dict(self) -> dict[str, Any]:
        payload = asdict(self)
        payload["sheets"] = [sheet.to_dict() for sheet in self.sheets]
        return payload


@dataclass(slots=True)
class GeneratedReportSpec:
    format: str = "xlsx"
    file_name: str = "report"
    title: str = "Report"
    summary: str = ""
    sections: list[dict[str, str]] = field(default_factory=list)
    tables: list[GeneratedTableSpec] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


def _normalize_text(value: object) -> str:
    return str(value or "").strip()


def _normalize_format(value: object) -> str:
    normalized = _normalize_text(value).lower().lstrip(".")
    aliases = {
        "excel": "xlsx",
        "xls": "xlsx",
        "spreadsheet": "xlsx",
        "word": "docx",
        "document": "docx",
        "markdown": "md",
        "text": "txt",
    }
    normalized = aliases.get(normalized, normalized)
    if normalized not in SUPPORTED_GENERATED_FILE_FORMATS:
        raise GeneratedFileError(
            f"Unsupported generated file format: {normalized or 'empty'}",
            error_code="unsupported_format",
            field_path="format",
            suggested_fix="Use one of: xlsx, csv, docx, pdf, txt, md, json.",
        )
    return normalized


def _normalize_file_name(file_format: str, index: int, suggested_name: object = "", fallback_title: object = "") -> str:
    raw_name = _normalize_text(suggested_name) or _normalize_text(fallback_title)
    if not raw_name:
        raw_name = f"generated-file-{index}.{file_format}"
    name = Path(raw_name.replace("\\", "/")).name
    name = re.sub(r"[\x00-\x1f<>:\"/\\|?*]+", "_", name).strip(" ._")
    if not name:
        name = f"generated-file-{index}.{file_format}"
    suffix = f".{file_format}"
    if file_format == "md":
        suffix = ".md"
    if not name.lower().endswith(suffix):
        name = f"{name}{suffix}"
    return name[:180]


def _sanitize_excel_sheet_title(raw_title: object, *, index: int, used_titles: set[str]) -> str:
    fallback = f"Sheet {index}"
    title = _normalize_text(raw_title) or fallback
    title = _EXCEL_INVALID_SHEET_TITLE_RE.sub("_", title).strip(" .")
    if not title:
        title = fallback
    base = title[:31].strip(" .") or fallback[:31]
    candidate = base
    suffix_index = 2
    while candidate.lower() in used_titles:
        suffix = f"_{suffix_index}"
        max_base_len = max(1, 31 - len(suffix))
        candidate_base = base[:max_base_len].strip(" .") or fallback[:max_base_len]
        candidate = f"{candidate_base}{suffix}"
        suffix_index += 1
    used_titles.add(candidate.lower())
    return candidate


def _normalize_column_key(value: object) -> str:
    text = _normalize_text(value).casefold().replace("\u0451", "\u0435")
    text = re.sub(r"[\s._-]+", " ", text)
    return text.strip(" :")


def _normalize_table_columns(columns: object) -> list[dict[str, str]]:
    if not isinstance(columns, list):
        return []
    result: list[dict[str, str]] = []
    for index, item in enumerate(columns, start=1):
        if isinstance(item, dict):
            key = _normalize_text(item.get("key") or item.get("id") or item.get("name") or item.get("field"))
            label = _normalize_text(item.get("label") or item.get("title") or item.get("header") or key)
        else:
            key = _normalize_text(item)
            label = key
        if not key and not label:
            continue
        result.append({"key": key or label or f"column_{index}", "label": label or key or f"Column {index}"})
    return result


def _equipment_canonical_for_key(key: object) -> str | None:
    return _EQUIPMENT_ALIAS_LOOKUP.get(_normalize_column_key(key))


def _is_blank_row_payload(row: object) -> bool:
    if isinstance(row, dict):
        return not any(_normalize_text(value) for value in row.values())
    if isinstance(row, list):
        return not any(_normalize_text(value) for value in row)
    return not bool(_normalize_text(row))


def _looks_like_equipment_headers(headers: list[object]) -> bool:
    canonical_keys = {
        canonical
        for header in headers
        for canonical in [_equipment_canonical_for_key(header)]
        if canonical
    }
    return "inv_no" in canonical_keys and bool(
        canonical_keys.intersection({"model_name", "type_name", "owner_name", "status", "branch", "location", "serial_no"})
    )


def _equipment_header_index(headers: list[str], canonical: str) -> int | None:
    for index, header in enumerate(headers):
        if _equipment_canonical_for_key(header) == canonical:
            return index
    return None


def _ensure_equipment_serial_column(table: list[list[str]]) -> list[list[str]]:
    if not table:
        return table
    headers = list(table[0] or [])
    if not _looks_like_equipment_headers(headers):
        return table
    inv_index = _equipment_header_index(headers, "inv_no")
    serial_index = _equipment_header_index(headers, "serial_no")
    if inv_index is None:
        return table
    result = [list(row or []) for row in table]
    if serial_index is not None:
        for row in result[1:]:
            while len(row) <= serial_index:
                row.append("")
            if not _normalize_text(row[serial_index]):
                row[serial_index] = _EMPTY_EQUIPMENT_SERIAL
        return result

    insert_at = inv_index + 1
    result[0].insert(insert_at, "\u0421\u0435\u0440\u0438\u0439\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440")
    for row in result[1:]:
        while len(row) < insert_at:
            row.append("")
        row.insert(insert_at, _EMPTY_EQUIPMENT_SERIAL)
    return result


def _normalize_dict_rows(rows: list[dict[str, Any]], *, columns: list[dict[str, str]] | None = None) -> list[list[str]]:
    keys: list[str] = []
    labels: list[str] = []
    for column in list(columns or []):
        key = _normalize_text(column.get("key"))
        label = _normalize_text(column.get("label")) or key
        if key and key not in keys:
            keys.append(key)
            labels.append(label)
    if not keys:
        for row in rows:
            for key in (row or {}).keys():
                normalized_key = _normalize_text(key)
                if normalized_key and normalized_key not in keys:
                    keys.append(normalized_key)
                    labels.append(normalized_key)
    table = [labels or keys] + [[_stringify_cell((row or {}).get(key)) for key in keys] for row in rows]
    return _ensure_equipment_serial_column(table)


def _normalize_list_rows(rows: list[list[Any]]) -> list[list[str]]:
    table = [[_stringify_cell(cell) for cell in row] for row in rows]
    return _ensure_equipment_serial_column(table)


def _recover_flat_equipment_table(rows: list[object]) -> list[list[str]]:
    cells = [_stringify_cell(row) for row in rows]
    header_count = 0
    for cell in cells:
        if _equipment_canonical_for_key(cell):
            header_count += 1
            continue
        break
    headers = cells[:header_count]
    if header_count < 2 or not _looks_like_equipment_headers(headers) or len(cells) % header_count != 0:
        raise GeneratedFileError(
            "Generated file rows must be a table: use an array of row arrays or row objects, not a flat list.",
            error_code="invalid_rows_shape",
            field_path="rows",
            suggested_fix="Pass rows as [[header...], [row...]] or as [{column: value}].",
        )
    table = [cells[index:index + header_count] for index in range(0, len(cells), header_count)]
    return _ensure_equipment_serial_column(table)


def _recover_flat_rows_with_columns(rows: list[object], columns: list[dict[str, str]]) -> list[list[str]]:
    labels = [_normalize_text(column.get("label")) or _normalize_text(column.get("key")) for column in columns]
    labels = [label for label in labels if label]
    if not labels:
        return _recover_flat_equipment_table(rows)
    cells = [_stringify_cell(row) for row in rows]
    column_count = len(labels)
    if len(cells) % column_count != 0:
        raise GeneratedFileError(
            "Generated file flat rows do not match the provided columns.",
            error_code="invalid_rows_shape",
            field_path="rows",
            suggested_fix="Make rows a list of row arrays or provide a cell count divisible by columns count.",
        )
    table = [labels]
    table.extend(cells[index:index + column_count] for index in range(0, len(cells), column_count))
    return _ensure_equipment_serial_column(table)


def _stringify_cell(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _normalize_rows(
    rows: object,
    *,
    columns: object = None,
    limit: int = MAX_GENERATED_FILE_ROWS_PER_SHEET,
    field_path: str = "rows",
) -> list[list[str]]:
    if not isinstance(rows, list):
        return []
    source_rows = [row for row in rows if not _is_blank_row_payload(row)]
    if not source_rows:
        return []
    normalized_columns = _normalize_table_columns(columns)
    if all(isinstance(row, dict) for row in source_rows):
        dict_rows = [row for row in source_rows if isinstance(row, dict)]
        result = _normalize_dict_rows(dict_rows, columns=normalized_columns)
    elif all(isinstance(row, list) for row in source_rows):
        list_rows = [row for row in source_rows if isinstance(row, list)]
        result = _normalize_list_rows(list_rows)
    elif all(not isinstance(row, (dict, list)) for row in source_rows):
        result = _recover_flat_rows_with_columns(source_rows, normalized_columns)
    else:
        raise GeneratedFileError(
            "Generated file rows must use one table shape: row arrays or row objects.",
            error_code="mixed_rows_shape",
            field_path=field_path,
            suggested_fix="Use either row arrays for every row, or row objects for every row.",
        )
    if len(result) > limit:
        raise GeneratedFileError(
            f"Generated file sheet has more than {limit} rows",
            error_code="row_limit_exceeded",
            field_path=field_path,
            suggested_fix=f"Filter the data or split it so each sheet has at most {limit} rows.",
        )
    return result


def _normalize_sheets(sheets: object) -> list[dict[str, Any]]:
    if not isinstance(sheets, list):
        return []
    result: list[dict[str, Any]] = []
    used_titles: set[str] = set()
    for index, sheet in enumerate(sheets, start=1):
        payload = sheet if isinstance(sheet, dict) else {}
        rows = _normalize_rows(
            payload.get("rows"),
            columns=payload.get("columns"),
            field_path=f"sheets[{index - 1}].rows",
        )
        if not rows:
            continue
        result.append(
            {
                "title": _sanitize_excel_sheet_title(payload.get("title"), index=index, used_titles=used_titles),
                "rows": rows,
                "header_row_index": int(payload.get("header_row_index") or 1),
            }
        )
    return result


def _normalize_generated_file_spec(item: dict[str, Any], *, index: int) -> GeneratedFileSpec:
    file_format = _normalize_format(item.get("format") or item.get("kind") or "xlsx")
    title = _normalize_text(item.get("title"))
    rows = _normalize_rows(item.get("rows"), columns=item.get("columns"), field_path="rows")
    sheets = [
        GeneratedSheetSpec(
            title=_normalize_text(sheet.get("title")),
            rows=list(sheet.get("rows") or []),
            header_row_index=int(sheet.get("header_row_index") or 1),
        )
        for sheet in _normalize_sheets(item.get("sheets"))
    ]
    spec = GeneratedFileSpec(
        format=file_format,
        file_name=_normalize_file_name(file_format, index, item.get("file_name"), title),
        title=title,
        content=item.get("content"),
        rows=rows,
        sheets=sheets,
        metadata=item.get("metadata") if isinstance(item.get("metadata"), dict) else {},
    )
    payload, _ = _generated_file_payload(spec.to_dict())
    size_bytes = len(payload)
    if size_bytes > MAX_GENERATED_FILE_BYTES:
        raise GeneratedFileError(
            f"Generated file '{spec.file_name}' is larger than {MAX_GENERATED_FILE_BYTES} bytes",
            error_code="file_size_limit_exceeded",
            field_path=f"files[{index - 1}]",
            suggested_fix="Reduce rows/content or create a smaller file.",
        )
    spec.size_bytes = size_bytes
    return spec


def normalize_generated_file_specs(files: list[dict[str, Any]] | None) -> list[dict[str, Any]]:
    source_files = [item for item in list(files or []) if isinstance(item, dict)]
    if len(source_files) > MAX_GENERATED_FILES_PER_RUN:
        raise GeneratedFileError(
            f"Generated files limit is {MAX_GENERATED_FILES_PER_RUN} files per run",
            error_code="file_count_limit_exceeded",
            field_path="files",
            suggested_fix=f"Create at most {MAX_GENERATED_FILES_PER_RUN} files in one AI run.",
        )
    specs: list[dict[str, Any]] = []
    for index, item in enumerate(source_files, start=1):
        specs.append(_normalize_generated_file_spec(item, index=index).to_dict())
    return specs


def _to_upload_file(file_name: str, payload: bytes, content_type: str) -> UploadFile:
    file_obj = io.BytesIO(bytes(payload or b""))
    return UploadFile(filename=file_name, file=file_obj, headers={"content-type": content_type})


def _text_from_spec(item: dict[str, Any]) -> str:
    content = item.get("content")
    if content is None and item.get("rows"):
        return "\n".join([", ".join(row) for row in list(item.get("rows") or [])])
    if isinstance(content, (dict, list)):
        return json.dumps(content, ensure_ascii=False, indent=2)
    return str(content or "")


def _report_metadata(item: dict[str, Any]) -> dict[str, Any] | None:
    metadata = item.get("metadata") if isinstance(item.get("metadata"), dict) else {}
    if not bool(metadata.get("report")):
        return None
    return metadata


def _report_tables_from_metadata(metadata: dict[str, Any]) -> list[dict[str, Any]]:
    tables: list[dict[str, Any]] = []
    for table in list(metadata.get("report_tables") or []):
        if not isinstance(table, dict):
            continue
        rows = _normalize_rows(
            table.get("rows"),
            columns=table.get("columns"),
            field_path="metadata.report_tables.rows",
        )
        if not rows:
            continue
        tables.append({"title": _normalize_text(table.get("title")) or "Data", "rows": rows})
    return tables


def _report_sections_from_metadata(metadata: dict[str, Any]) -> list[dict[str, str]]:
    sections: list[dict[str, str]] = []
    for section in list(metadata.get("report_sections") or []):
        if not isinstance(section, dict):
            continue
        heading = _normalize_text(section.get("heading")) or "Section"
        body = _normalize_text(section.get("body"))
        if heading or body:
            sections.append({"heading": heading, "body": body})
    return sections


def _csv_payload(item: dict[str, Any]) -> bytes:
    output = io.StringIO()
    writer = csv.writer(output)
    rows = item.get("rows")
    if isinstance(rows, list) and rows:
        for row in rows:
            writer.writerow([_stringify_cell(cell) for cell in list(row or [])])
    else:
        output.write(_text_from_spec(item))
    return output.getvalue().encode("utf-8-sig")


def _xlsx_payload(item: dict[str, Any]) -> bytes:
    if Workbook is None:
        raise GeneratedFileError("openpyxl is not installed")
    workbook = Workbook()
    workbook.remove(workbook.active)
    sheets = item.get("sheets")
    used_titles: set[str] = set()
    if isinstance(sheets, list) and sheets:
        for index, sheet_payload in enumerate(sheets, start=1):
            title = _sanitize_excel_sheet_title((sheet_payload or {}).get("title"), index=index, used_titles=used_titles)
            sheet = workbook.create_sheet(title=title)
            for row in list((sheet_payload or {}).get("rows") or []):
                sheet.append([_stringify_cell(cell) for cell in list(row or [])])
            _style_xlsx_sheet(sheet, header_row_index=int((sheet_payload or {}).get("header_row_index") or 1))
    else:
        title = _sanitize_excel_sheet_title(item.get("title"), index=1, used_titles=used_titles)
        sheet = workbook.create_sheet(title=title)
        rows = item.get("rows")
        if isinstance(rows, list) and rows:
            for row in rows:
                sheet.append([_stringify_cell(cell) for cell in list(row or [])])
        else:
            sheet.append([_text_from_spec(item)])
        _style_xlsx_sheet(sheet, header_row_index=1)
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _style_xlsx_sheet(sheet, *, header_row_index: int = 1) -> None:
    if Font is None or PatternFill is None or Alignment is None or Border is None or Side is None or get_column_letter is None:
        return
    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(color="FFFFFF", bold=True)
    title_font = Font(size=14, bold=True, color="1F4E79")
    thin = Side(style="thin", color="D9E2F3")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in sheet.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border
    if sheet.max_row >= 1:
        for cell in sheet[1]:
            cell.font = title_font if header_row_index != 1 else header_font
            if header_row_index == 1:
                cell.fill = header_fill
    if 1 <= header_row_index <= sheet.max_row:
        for cell in sheet[header_row_index]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        sheet.freeze_panes = f"A{header_row_index + 1}"
        sheet.auto_filter.ref = sheet.dimensions
    for column_cells in sheet.columns:
        width = 12
        column_letter = get_column_letter(column_cells[0].column)
        for cell in column_cells:
            width = max(width, min(42, len(_normalize_text(cell.value)) + 2))
        sheet.column_dimensions[column_letter].width = width


def _docx_payload(item: dict[str, Any]) -> bytes:
    if Document is None:
        raise GeneratedFileError("python-docx is not installed")
    document = Document()
    title = _normalize_text(item.get("title"))
    if title:
        document.add_heading(title, level=1)
    report_meta = _report_metadata(item)
    if report_meta:
        summary = _normalize_text(report_meta.get("report_summary"))
        if summary:
            document.add_paragraph(summary)
        for section in _report_sections_from_metadata(report_meta):
            document.add_heading(section["heading"], level=2)
            if section["body"]:
                document.add_paragraph(section["body"])
        for table_payload in _report_tables_from_metadata(report_meta):
            document.add_heading(table_payload["title"], level=2)
            _append_docx_table(document, list(table_payload.get("rows") or []))
    else:
        text = _text_from_spec(item)
        for line in text.splitlines() or [""]:
            document.add_paragraph(line)
    if item.get("rows"):
        rows = list(item.get("rows") or [])
        if rows:
            _append_docx_table(document, rows)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _append_docx_table(document: Any, rows: list[list[Any]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows) or 1
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for cell_index in range(column_count):
            value = row[cell_index] if cell_index < len(row) else ""
            table.rows[row_index].cells[cell_index].text = _stringify_cell(value)
            if row_index == 0:
                for paragraph in table.rows[row_index].cells[cell_index].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _pdf_font_path() -> str | None:
    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/ARIAL.TTF"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return str(path)
    return None


def _register_pdf_font() -> str:
    global _PDF_FONT_REGISTERED
    if canvas is None or A4 is None or pdfmetrics is None or TTFont is None:
        raise GeneratedFileError("reportlab is not installed")
    if not _PDF_FONT_REGISTERED:
        font_path = _pdf_font_path()
        if not font_path:
            raise GeneratedFileError("PDF font unavailable")
        pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME, font_path))
        _PDF_FONT_REGISTERED = True
    return _PDF_FONT_NAME


def _pdf_payload(item: dict[str, Any]) -> bytes:
    font_name = _register_pdf_font()
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 48
    title = _normalize_text(item.get("title"))

    def new_page(font_size: int = 10) -> None:
        nonlocal y
        pdf.showPage()
        pdf.setFont(font_name, font_size)
        y = height - 48

    def draw_wrapped(text: object, *, font_size: int = 10, line_height: int = 14, max_chars: int = 108) -> None:
        nonlocal y
        pdf.setFont(font_name, font_size)
        raw_lines = _normalize_text(text).splitlines() or [""]
        wrapped_lines: list[str] = []
        for raw_line in raw_lines:
            line = raw_line
            if not line:
                wrapped_lines.append("")
                continue
            while len(line) > max_chars:
                wrapped_lines.append(line[:max_chars])
                line = line[max_chars:]
            wrapped_lines.append(line)
        for line in wrapped_lines:
            if y < 48:
                new_page(font_size)
            pdf.drawString(40, y, line)
            y -= line_height

    def draw_table(rows: list[list[Any]]) -> None:
        nonlocal y
        if not rows:
            return
        column_count = max(len(row) for row in rows) or 1
        usable_width = width - 80
        column_width = usable_width / column_count
        row_height = 18
        pdf.setFont(font_name, 8)
        for row_index, row in enumerate(rows):
            if y < 64:
                new_page(8)
            x = 40
            for cell_index in range(column_count):
                value = _stringify_cell(row[cell_index] if cell_index < len(row) else "")
                pdf.rect(x, y - row_height + 4, column_width, row_height, stroke=1, fill=0)
                if row_index == 0:
                    pdf.setFont(font_name, 8)
                pdf.drawString(x + 3, y - 9, value[: max(8, int(column_width / 4))])
                x += column_width
            y -= row_height
        y -= 8

    if title:
        draw_wrapped(title[:160], font_size=15, line_height=20, max_chars=90)
        y -= 8
    report_meta = _report_metadata(item)
    if report_meta:
        summary = _normalize_text(report_meta.get("report_summary"))
        if summary:
            draw_wrapped(summary, font_size=10)
            y -= 8
        for section in _report_sections_from_metadata(report_meta):
            draw_wrapped(section["heading"], font_size=12, line_height=16, max_chars=90)
            if section["body"]:
                draw_wrapped(section["body"], font_size=10)
            y -= 6
        for table_payload in _report_tables_from_metadata(report_meta):
            draw_wrapped(table_payload["title"], font_size=12, line_height=16, max_chars=90)
            draw_table(list(table_payload.get("rows") or []))
    elif item.get("rows"):
        draw_table(list(item.get("rows") or []))
    else:
        draw_wrapped(_text_from_spec(item), font_size=11, line_height=16)
    pdf.save()
    return buffer.getvalue()


def _json_payload(item: dict[str, Any]) -> bytes:
    content = item.get("content")
    if content is None:
        content = {"title": item.get("title"), "rows": item.get("rows"), "sheets": item.get("sheets")}
    if isinstance(content, str):
        try:
            content = json.loads(content)
        except Exception:
            pass
    return json.dumps(content, ensure_ascii=False, indent=2).encode("utf-8")


def _generated_file_payload(item: dict[str, Any]) -> tuple[bytes, str]:
    file_format = _normalize_format(item.get("format") or item.get("kind"))
    if file_format == "csv":
        payload = _csv_payload(item)
    elif file_format == "xlsx":
        payload = _xlsx_payload(item)
    elif file_format == "docx":
        payload = _docx_payload(item)
    elif file_format == "pdf":
        payload = _pdf_payload(item)
    elif file_format == "json":
        payload = _json_payload(item)
    else:
        payload = _text_from_spec(item).encode("utf-8")
    return payload, _CONTENT_TYPES[file_format]


def build_generated_uploads(artifacts: list[dict[str, Any]] | None) -> list[UploadFile]:
    uploads: list[UploadFile] = []
    for item in normalize_generated_file_specs(list(artifacts or [])):
        payload, content_type = _generated_file_payload(item)
        uploads.append(_to_upload_file(str(item["file_name"]), payload, content_type))
    return uploads
