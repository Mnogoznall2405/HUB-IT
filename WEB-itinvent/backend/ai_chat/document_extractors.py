from __future__ import annotations

import csv
import html
import io
import json
import mimetypes
import re
import zipfile
from email import policy
from email.parser import BytesParser
from html.parser import HTMLParser
from pathlib import Path
from typing import Callable
from xml.etree import ElementTree as ET

try:
    import fitz  # type: ignore
except Exception:  # pragma: no cover
    fitz = None

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None

try:
    from openpyxl import load_workbook  # type: ignore
except Exception:  # pragma: no cover
    load_workbook = None

try:
    from pypdf import PdfReader  # type: ignore
except Exception:  # pragma: no cover
    PdfReader = None


TextExtractor = Callable[[bytes, str], str]
PDF_OCR_MAX_PAGES = 3
ZIP_IMAGE_OCR_MAX_FILES = 3
_INLINE_WHITESPACE_RE = re.compile(r"[ \t\f\v]+")
_BINARY_LATIN_TEXT_RE = re.compile(rb"[\x20-\x7E\xA0-\xFF]{4,}")
_BINARY_UTF16LE_TEXT_RE = re.compile(rb"(?:(?:[\x20-\x7E]\x00)|(?:[\xA0-\xFF]\x00)|(?:[\x00-\xFF]\x04)){4,}")
_RTF_UNICODE_RE = re.compile(r"\\u(-?\d+)\??")
_RTF_HEX_RE = re.compile(r"\\'([0-9a-fA-F]{2})")
_RTF_DESTINATION_RE = re.compile(r"\{\\\*[^{}]*\}")
_RTF_CONTROL_WORD_RE = re.compile(r"\\[a-zA-Z]+-?\d* ?")
_RTF_KNOWN_GROUP_RE = re.compile(r"\{\\(?:fonttbl|colortbl|stylesheet|info)(?:[^{}]|\{[^{}]*\})*\}", flags=re.DOTALL)
_OOXML_WORD_EXTENSIONS = {".docx", ".docm", ".dotx", ".dotm"}
_OOXML_SHEET_EXTENSIONS = {".xlsx", ".xlsm", ".xltx", ".xltm"}
_OOXML_PRESENTATION_EXTENSIONS = {".pptx", ".pptm", ".potx", ".potm", ".ppsx", ".ppsm"}
_ODF_EXTENSIONS = {".odt", ".ods", ".odp"}
_LEGACY_BINARY_DOCUMENT_EXTENSIONS = {".doc", ".xls", ".ppt"}
_TEXT_EXTENSIONS = {
    ".txt", ".md", ".log", ".xml", ".yml", ".yaml", ".ini", ".cfg", ".conf", ".sql",
}


def _read_bytes(path: Path) -> bytes:
    return path.read_bytes()


def _decode_text(payload: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return payload.decode(encoding)
        except Exception:
            continue
    return payload.decode("utf-8", errors="ignore")


def _collapse_inline_whitespace(value: object) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    return _INLINE_WHITESPACE_RE.sub(" ", text).strip()


def _finalize_text(parts: list[str]) -> str:
    lines: list[str] = []
    for part in list(parts or []):
        raw = str(part or "").replace("\r\n", "\n").replace("\r", "\n")
        for line in raw.split("\n"):
            normalized = _collapse_inline_whitespace(line)
            if normalized:
                lines.append(normalized)
            elif lines and lines[-1] != "":
                lines.append("")
    while lines and lines[0] == "":
        lines.pop(0)
    while lines and lines[-1] == "":
        lines.pop()
    return "\n".join(lines).strip()


def _dedupe_preserving_order(parts: list[str]) -> list[str]:
    seen: set[str] = set()
    unique: list[str] = []
    for part in list(parts or []):
        normalized = _collapse_inline_whitespace(part).lower()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        unique.append(str(part))
    return unique


def _looks_like_meaningful_text(value: object) -> bool:
    text = _collapse_inline_whitespace(value)
    if len(text) < 4:
        return False
    meaningful = sum(1 for char in text if char.isalnum())
    return meaningful >= max(3, len(text) // 4)


def _natural_sort_key(value: str) -> list[object]:
    return [
        int(token) if token.isdigit() else token.lower()
        for token in re.split(r"(\d+)", str(value or ""))
    ]


def _run_image_extractor(
    image_extractor: TextExtractor | None,
    *,
    payload: bytes,
    mime_type: str,
    prompt: str = "",
) -> str:
    if image_extractor is None or not payload:
        return ""
    try:
        return str(
            image_extractor(
                image_bytes=payload,
                mime_type=str(mime_type or "image/png").strip() or "image/png",
                prompt=str(prompt or "").strip(),
            ) or ""
        ).strip()
    except TypeError:
        try:
            return str(image_extractor(payload, mime_type) or "").strip()
        except Exception:
            return ""
    except Exception:
        return ""


def _extract_xml_text(payload: bytes) -> str:
    try:
        root = ET.fromstring(payload)
    except Exception:
        return ""
    parts: list[str] = []
    for element in root.iter():
        text = _collapse_inline_whitespace(element.text)
        if text:
            parts.append(text)
    return _finalize_text(parts)


def _extract_zip_image_ocr(
    path: Path,
    *,
    member_prefixes: tuple[str, ...],
    image_extractor: TextExtractor | None = None,
) -> str:
    if image_extractor is None:
        return ""
    try:
        with zipfile.ZipFile(path) as archive:
            names = sorted(archive.namelist(), key=_natural_sort_key)
            ocr_parts: list[str] = []
            for name in names:
                if not any(name.startswith(prefix) for prefix in member_prefixes):
                    continue
                mime_type = str(mimetypes.guess_type(name)[0] or "").strip().lower()
                if not mime_type.startswith("image/"):
                    continue
                try:
                    payload = archive.read(name)
                except Exception:
                    continue
                extracted = _run_image_extractor(
                    image_extractor,
                    payload=payload,
                    mime_type=mime_type or "image/png",
                    prompt="Extract the readable text from this document image. Return plain text only.",
                )
                if extracted:
                    ocr_parts.append(extracted)
                if len(ocr_parts) >= ZIP_IMAGE_OCR_MAX_FILES:
                    break
            return _finalize_text(ocr_parts)
    except Exception:
        return ""


def _extract_binary_strings(payload: bytes) -> str:
    parts: list[str] = []
    for match in _BINARY_UTF16LE_TEXT_RE.findall(payload):
        try:
            candidate = match.decode("utf-16le", errors="ignore")
        except Exception:
            continue
        candidate = re.sub(r"^[0-9]+(?=[А-Яа-яЁё]{3,})", "", candidate)
        if _looks_like_meaningful_text(candidate):
            parts.append(candidate)
    for match in _BINARY_LATIN_TEXT_RE.findall(payload):
        candidate = _decode_text(match)
        if _looks_like_meaningful_text(candidate):
            parts.append(candidate)
    try:
        utf16_text = payload.decode("utf-16le", errors="ignore")
    except Exception:
        utf16_text = ""
    if utf16_text:
        for candidate in re.findall(r"[A-Za-zА-Яа-яЁё][0-9A-Za-zА-Яа-яЁё .,:;()/_-]{3,}", utf16_text):
            if _looks_like_meaningful_text(candidate):
                parts.append(candidate)
    return _finalize_text(_dedupe_preserving_order(parts))


def _extract_pdf_text(path: Path, *, image_extractor: TextExtractor | None = None) -> str:
    if PdfReader is not None:
        try:
            reader = PdfReader(str(path))
            parts = [str(page.extract_text() or "").strip() for page in reader.pages]
            text = "\n\n".join(part for part in parts if part)
            if text.strip():
                return text
        except Exception:
            pass
    if fitz is not None:
        try:
            doc = fitz.open(str(path))
            try:
                parts = [str(page.get_text("text") or "").strip() for page in doc]
                text = "\n\n".join(part for part in parts if part)
                if text.strip():
                    return text
                if image_extractor is not None:
                    ocr_parts: list[str] = []
                    pages_limit = min(max(len(doc), 0), max(1, int(PDF_OCR_MAX_PAGES)))
                    matrix = fitz.Matrix(2.0, 2.0)
                    for page_index in range(pages_limit):
                        page = doc.load_page(page_index)
                        pix = page.get_pixmap(matrix=matrix, alpha=False)
                        extracted = _run_image_extractor(
                            image_extractor,
                            payload=pix.tobytes("png"),
                            mime_type="image/png",
                            prompt="Extract the readable text from this PDF page. Return plain text only.",
                        )
                        if extracted:
                            ocr_parts.append(extracted)
                    return "\n\n".join(part for part in ocr_parts if part)
            finally:
                doc.close()
        except Exception:
            return ""
    return ""


def _extract_docx_text(path: Path, *, image_extractor: TextExtractor | None = None) -> str:
    parts: list[str] = []
    if Document is not None:
        try:
            document = Document(str(path))
            for paragraph in document.paragraphs:
                text = _collapse_inline_whitespace(paragraph.text)
                if text:
                    parts.append(text)
            for table in document.tables:
                for row in table.rows:
                    cells = [
                        _collapse_inline_whitespace(cell.text)
                        for cell in list(row.cells or [])
                    ]
                    cells = [cell for cell in cells if cell]
                    if cells:
                        parts.append(" | ".join(cells))
            for section in document.sections:
                for paragraph in section.header.paragraphs:
                    text = _collapse_inline_whitespace(paragraph.text)
                    if text:
                        parts.append(text)
                for paragraph in section.footer.paragraphs:
                    text = _collapse_inline_whitespace(paragraph.text)
                    if text:
                        parts.append(text)
        except Exception:
            pass
    try:
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            xml_names: list[str] = []
            if "word/document.xml" in names:
                xml_names.append("word/document.xml")
            for prefix in ("word/header", "word/footer", "word/comments", "word/footnotes", "word/endnotes"):
                xml_names.extend(
                    name
                    for name in sorted(names, key=_natural_sort_key)
                    if name.startswith(prefix) and name.endswith(".xml")
                )
            for name in xml_names:
                try:
                    text = _extract_xml_text(archive.read(name))
                except Exception:
                    continue
                if text:
                    parts.append(text)
    except Exception:
        pass
    text = _finalize_text(_dedupe_preserving_order(parts))
    if text:
        return text
    return _extract_zip_image_ocr(path, member_prefixes=("word/media/",), image_extractor=image_extractor)


def _extract_pptx_text(path: Path, *, image_extractor: TextExtractor | None = None) -> str:
    parts: list[str] = []
    try:
        with zipfile.ZipFile(path) as archive:
            names = sorted(archive.namelist(), key=_natural_sort_key)
            slide_names = [
                name
                for name in names
                if name.startswith("ppt/slides/slide") and name.endswith(".xml")
            ]
            for index, name in enumerate(slide_names, start=1):
                text = _extract_xml_text(archive.read(name))
                if text:
                    parts.append(f"# Slide {index}\n{text}")
            note_names = [
                name
                for name in names
                if name.startswith("ppt/notesSlides/notesSlide") and name.endswith(".xml")
            ]
            for index, name in enumerate(note_names, start=1):
                text = _extract_xml_text(archive.read(name))
                if text:
                    parts.append(f"# Slide notes {index}\n{text}")
    except Exception:
        return ""
    text = _finalize_text(parts)
    if text:
        return text
    return _extract_zip_image_ocr(path, member_prefixes=("ppt/media/",), image_extractor=image_extractor)


def _extract_xlsx_text(path: Path) -> str:
    if load_workbook is not None:
        try:
            workbook = load_workbook(str(path), read_only=True, data_only=True)
        except Exception:
            workbook = None
        if workbook is not None:
            try:
                lines: list[str] = []
                for sheet in workbook.worksheets:
                    lines.append(f"# Sheet: {sheet.title}")
                    for row in sheet.iter_rows(values_only=True):
                        cells = [str(cell).strip() for cell in row if cell not in (None, "")]
                        if cells:
                            lines.append(" | ".join(cells))
                    lines.append("")
                text = "\n".join(lines).strip()
                if text:
                    return text
            finally:
                try:
                    workbook.close()
                except Exception:
                    pass
    try:
        with zipfile.ZipFile(path) as archive:
            parts: list[str] = []
            for name in sorted(archive.namelist(), key=_natural_sort_key):
                if not name.startswith("xl/") or not name.endswith(".xml"):
                    continue
                if name.startswith("xl/theme/") or name.startswith("xl/styles"):
                    continue
                text = _extract_xml_text(archive.read(name))
                if text:
                    parts.append(text)
            return _finalize_text(parts)
    except Exception:
        return ""


def _extract_odf_text(path: Path, *, image_extractor: TextExtractor | None = None) -> str:
    try:
        with zipfile.ZipFile(path) as archive:
            parts: list[str] = []
            for name in ("content.xml", "meta.xml"):
                if name not in archive.namelist():
                    continue
                text = _extract_xml_text(archive.read(name))
                if text:
                    parts.append(text)
            text = _finalize_text(parts)
            if text:
                return text
    except Exception:
        return ""
    return _extract_zip_image_ocr(path, member_prefixes=("Pictures/",), image_extractor=image_extractor)


def _extract_delimited_text(path: Path, *, delimiter: str = ",") -> str:
    try:
        payload = _read_bytes(path)
        decoded = _decode_text(payload)
        reader = csv.reader(io.StringIO(decoded), delimiter=str(delimiter or ","))
        return "\n".join(" | ".join(str(cell or "").strip() for cell in row) for row in reader)
    except Exception:
        return ""


def _extract_json_text(path: Path) -> str:
    try:
        parsed = json.loads(_decode_text(_read_bytes(path)))
        return json.dumps(parsed, ensure_ascii=False, indent=2)
    except Exception:
        return _decode_text(_read_bytes(path))


class _HtmlTextParser(HTMLParser):
    _BLOCK_TAGS = {
        "address", "article", "aside", "blockquote", "div", "dl", "fieldset", "figcaption",
        "figure", "footer", "form", "h1", "h2", "h3", "h4", "h5", "h6", "header", "li",
        "main", "nav", "ol", "p", "pre", "section", "table", "tr", "ul",
    }

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._parts: list[str] = []

    def handle_starttag(self, tag: str, attrs) -> None:  # type: ignore[override]
        if str(tag or "").lower() == "br":
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:  # type: ignore[override]
        if str(tag or "").lower() in self._BLOCK_TAGS:
            self._parts.append("\n")

    def handle_data(self, data: str) -> None:  # type: ignore[override]
        if data:
            self._parts.append(data)

    def get_text(self) -> str:
        return _finalize_text(["".join(self._parts)])


def _extract_html_text(path: Path) -> str:
    try:
        parser = _HtmlTextParser()
        parser.feed(_decode_text(_read_bytes(path)))
        parser.close()
        return parser.get_text()
    except Exception:
        return ""


def _extract_rtf_text(path: Path) -> str:
    try:
        raw = _read_bytes(path).decode("latin-1", errors="ignore")
    except Exception:
        return ""

    raw = _RTF_KNOWN_GROUP_RE.sub("", raw)
    raw = re.sub(r"\\pard(?:plain)?\b ?", "\n", raw)
    raw = re.sub(r"\\par\b ?", "\n", raw)
    raw = re.sub(r"\\line\b ?", "\n", raw)
    raw = re.sub(r"\\tab\b ?", "\t", raw)

    def replace_unicode(match: re.Match[str]) -> str:
        try:
            codepoint = int(match.group(1))
        except Exception:
            return ""
        if codepoint < 0:
            codepoint += 65536
        try:
            return chr(codepoint)
        except Exception:
            return ""

    def replace_hex(match: re.Match[str]) -> str:
        try:
            return bytes([int(match.group(1), 16)]).decode("cp1251", errors="ignore")
        except Exception:
            return ""

    text = _RTF_UNICODE_RE.sub(replace_unicode, raw)
    text = _RTF_HEX_RE.sub(replace_hex, text)
    while True:
        updated = _RTF_DESTINATION_RE.sub("", text)
        if updated == text:
            break
        text = updated
    text = text.replace("\\{", "{").replace("\\}", "}").replace("\\\\", "\\")
    text = _RTF_CONTROL_WORD_RE.sub("", text)
    text = text.replace("{", "").replace("}", "")
    return _finalize_text([html.unescape(text)])


def _extract_eml_text(path: Path) -> str:
    try:
        message = BytesParser(policy=policy.default).parsebytes(_read_bytes(path))
    except Exception:
        return ""

    parts: list[str] = []
    subject = _collapse_inline_whitespace(message.get("subject"))
    if subject:
        parts.append(f"Subject: {subject}")

    if message.is_multipart():
        for part in message.walk():
            content_type = str(part.get_content_type() or "").lower()
            if content_type not in {"text/plain", "text/html"}:
                continue
            try:
                payload = part.get_content()
            except Exception:
                continue
            if content_type == "text/html":
                parser = _HtmlTextParser()
                parser.feed(str(payload or ""))
                parser.close()
                extracted = parser.get_text()
            else:
                extracted = _finalize_text([str(payload or "")])
            if extracted:
                parts.append(extracted)
    else:
        try:
            payload = message.get_content()
        except Exception:
            payload = ""
        content_type = str(message.get_content_type() or "").lower()
        if content_type == "text/html":
            parser = _HtmlTextParser()
            parser.feed(str(payload or ""))
            parser.close()
            extracted = parser.get_text()
        else:
            extracted = _finalize_text([str(payload or "")])
        if extracted:
            parts.append(extracted)
    return _finalize_text(parts)


def extract_text_from_path(
    path: str | Path,
    *,
    file_name: str = "",
    mime_type: str = "",
    image_extractor: TextExtractor | None = None,
) -> str:
    target = Path(path)
    if not target.exists() or not target.is_file():
        return ""

    resolved_name = str(file_name or target.name).strip()
    normalized_suffix = Path(resolved_name).suffix.lower() or target.suffix.lower()
    normalized_mime_type = str(
        mime_type
        or mimetypes.guess_type(resolved_name)[0]
        or mimetypes.guess_type(target.name)[0]
        or ""
    ).strip().lower()

    if normalized_suffix in _TEXT_EXTENSIONS:
        return _decode_text(_read_bytes(target))
    if normalized_suffix in {".html", ".htm"}:
        return _extract_html_text(target)
    if normalized_suffix in {".json", ".jsonl"}:
        return _extract_json_text(target)
    if normalized_suffix == ".csv":
        return _extract_delimited_text(target, delimiter=",")
    if normalized_suffix == ".tsv":
        return _extract_delimited_text(target, delimiter="\t")
    if normalized_suffix == ".pdf":
        return _extract_pdf_text(target, image_extractor=image_extractor)
    if normalized_suffix in _OOXML_WORD_EXTENSIONS:
        return _extract_docx_text(target, image_extractor=image_extractor)
    if normalized_suffix in _OOXML_SHEET_EXTENSIONS:
        return _extract_xlsx_text(target)
    if normalized_suffix in _OOXML_PRESENTATION_EXTENSIONS:
        return _extract_pptx_text(target, image_extractor=image_extractor)
    if normalized_suffix in _ODF_EXTENSIONS:
        return _extract_odf_text(target, image_extractor=image_extractor)
    if normalized_suffix == ".rtf":
        return _extract_rtf_text(target)
    if normalized_suffix == ".eml":
        return _extract_eml_text(target)
    if normalized_suffix in _LEGACY_BINARY_DOCUMENT_EXTENSIONS:
        return _extract_binary_strings(_read_bytes(target))
    if normalized_mime_type.startswith("text/"):
        return _decode_text(_read_bytes(target))
    if normalized_mime_type.startswith("image/") and image_extractor is not None:
        return _run_image_extractor(
            image_extractor,
            payload=_read_bytes(target),
            mime_type=normalized_mime_type,
            prompt="Extract the full readable text from this image. Return plain text only.",
        )
    if normalized_mime_type in {
        "application/msword",
        "application/vnd.ms-excel",
        "application/vnd.ms-powerpoint",
    }:
        return _extract_binary_strings(_read_bytes(target))
    return ""
