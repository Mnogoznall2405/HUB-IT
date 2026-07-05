#!/usr/bin/env python3
"""Build HR onboarding portal guide DOCX for HUB-IT."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
ASSETS = ROOT / "documentation" / "user-guides" / "assets" / "hr-onboarding"
OUTPUT = ROOT / "documentation" / "user-guides" / "hr-onboarding-portal-guide.docx"

PRIMARY = RGBColor(0x19, 0x76, 0xD2)
MUTED = RGBColor(0x64, 0x74, 0x8B)


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shading)


def _style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(11)
    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Segoe UI"
        style.font.color.rgb = PRIMARY
        style.font.bold = True
        if level == 1:
            style.font.size = Pt(18)
        elif level == 2:
            style.font.size = Pt(14)
        else:
            style.font.size = Pt(12)


def _add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("HUB-IT | Портал трудоустройства | Конфиденциально")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def _add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(hdr_cells[i], "1976d2")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph("")


def _add_image(doc: Document, filename: str, caption: str, width_in: float = 6.5) -> None:
    path = ASSETS / filename
    if not path.exists():
        _add_paragraph(doc, f"[Изображение не найдено: {filename}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = MUTED
    doc.add_paragraph("")


def build_document() -> Document:
    doc = Document()
    _style_document(doc)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Портал сбора документов при трудоустройстве")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = PRIMARY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Правовая модель, процесс и интеграция в HUB-IT")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Внутренний документ | Версия 1.0 | {date.today().strftime('%d.%m.%Y')}\n"
        "Аудитория: HR, юристы, комплаенс"
    )
    meta_run.font.size = Pt(11)

    doc.add_page_break()

    doc.add_heading("Оглавление", level=1)
    toc_items = [
        "1. Назначение документа",
        "2. Правовая база",
        "3. Перечень документов при приеме на работу",
        "4. Процесс взаимодействия",
        "5. Согласия и локальные акты",
        "6. Роли и разграничение доступа",
        "7. Хранение и уничтожение данных",
        "8. Проверка подлинности документов",
        "9. Риски и антипаттерны",
        "Приложение A. Макеты экранов",
        "Приложение B. Интеграция в HUB-IT",
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    doc.add_heading("Дисклеймер", level=1)
    _add_paragraph(
        doc,
        "Настоящий документ подготовлен для внутреннего использования в HUB-IT и носит "
        "информационно-методический характер. Он не является юридической консультацией. "
        "Перед запуском портала необходимо согласование с корпоративным юристом и ответственным за ПДн.",
    )

    doc.add_heading("1. Назначение документа", level=1)
    _add_paragraph(
        doc,
        "Документ описывает законный порядок сбора документов от соискателей и принимаемых "
        "на работу сотрудников через электронный портал, а также то, как такой портал может "
        "быть встроен в платформу HUB-IT.",
    )
    _add_paragraph(doc, "Цели портала:", bold=True)
    _add_bullets(
        doc,
        [
            "Сократить бумажный обмен и ошибки при переносе данных в личное дело.",
            "Обеспечить конфиденциальность персональных данных (ПДн).",
            "Зафиксировать согласие субъекта ПДн до загрузки документов.",
            "Ограничить доступ к пакету документов уполномоченными сотрудниками.",
        ],
    )
    _add_paragraph(doc, "Границы применения:", bold=True)
    _add_bullets(
        doc,
        [
            "Соискатель - лицо, по которому еще не заключен трудовой договор.",
            "Принимаемый сотрудник - лицо, по которому принято решение о приеме.",
            "Портал не заменяет КЭДО, электронную трудовую книжку и архивное хранение личного дела.",
        ],
    )

    doc.add_heading("2. Правовая база", level=1)
    _add_table(
        doc,
        ["Норма", "Содержание для HR-портала"],
        [
            ["Ст. 65 ТК РФ", "Закрытый перечень документов при приеме на работу"],
            ["Ст. 86 ТК РФ", "Обязанности работодателя при обработке ПДн работников"],
            ["Ст. 88 ТК РФ", "Передача ПДн третьим лицам - с письменного согласия"],
            ["Ст. 22.2 ТК РФ", "Документы могут предъявляться в электронной форме"],
            ["Ст. 6 152-ФЗ", "Законные основания обработки ПДн"],
            ["Ст. 9 152-ФЗ", "Согласие отдельным документом (с 01.09.2025)"],
            ["Ст. 10 152-ФЗ", "Специальные категории ПДн (здоровье и др.)"],
            ["Ст. 10.1 152-ФЗ", "Распространение ПДн - отдельное согласие"],
            ["Ст. 11 152-ФЗ", "Биометрические ПДн - письменное согласие"],
            ["Ст. 19 152-ФЗ", "Меры защиты ПДн в информационных системах"],
            ["Ст. 21 152-ФЗ", "Уничтожение ПДн при достижении цели (30 дней)"],
            ["Ст. 22 152-ФЗ", "Уведомление Роскомнадзора об обработке ПДн"],
            ["ПП РФ № 1119", "Уровни защищенности ИСПДн"],
            ["420-ФЗ (2025)", "Ужесточенные штрафы за утечки ПДн"],
        ],
    )
    _add_paragraph(
        doc,
        "Для соискателей до трудоустройства основным основанием обработки является согласие "
        "субъекта ПДн. После заключения трудового договора часть обработки допускается без "
        "отдельного согласия в рамках исполнения трудового законодательства, но передача "
        "третьим лицам и распространение по-прежнему требуют отдельного правового основания.",
    )

    doc.add_heading("3. Перечень документов при приеме на работу", level=1)
    _add_paragraph(
        doc,
        "Работодатель вправе запросить только документы, предусмотренные ст. 65 ТК РФ и иными "
        "федеральными законами для конкретной должности.",
    )
    _add_table(
        doc,
        ["Документ", "Основание", "Срок хранения"],
        [
            ["Паспорт", "Ст. 65 ТК РФ", "Личное дело - 50 лет (архив)"],
            ["СНИЛС / регистрация в персучете", "Ст. 65 ТК РФ", "Личное дело - 50 лет"],
            ["Трудовая / сведения о стаже", "Ст. 65, 66.1 ТК РФ", "Личное дело - 50 лет"],
            ["Воинский учет", "Ст. 65 ТК РФ", "Личное дело - 50 лет"],
            ["Образование / квалификация", "Ст. 65 ТК РФ (если требуется)", "Личное дело - 50 лет"],
            ["Справка о судимости", "ФЗ для отдельных должностей", "По локальным правилам"],
            ["Пакет соискателя при отказе", "152-ФЗ", "Уничтожение в течение 30 дней"],
        ],
    )

    doc.add_heading("4. Процесс взаимодействия", level=1)
    _add_paragraph(doc, "Законная последовательность шагов:", bold=True)
    steps = [
        "HR создает карточку кандидата и отправляет персональную ссылку (email/SMS).",
        "Кандидат знакомится с политикой обработки ПДн.",
        "Кандидат подписывает отдельное согласие на обработку ПДн.",
        "Кандидат загружает документы по перечню ст. 65 ТК РФ.",
        "Система выполняет техническую проверку формата (OCR, MRZ).",
        "HR проверяет пакет и принимает решение.",
        "При приеме - перенос в личное дело/КЭДО; при отказе - уничтожение за 30 дней.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    _add_image(doc, "diagram-01-roles.png", "Рис. 1. Участники процесса и роли")
    _add_image(doc, "diagram-02-lifecycle.png", "Рис. 2. Жизненный цикл пакета документов")
    _add_image(doc, "diagram-03-consents.png", "Рис. 3. Документооборот согласий")
    _add_image(doc, "diagram-04-storage.png", "Рис. 4. Временный и постоянный контуры хранения")

    doc.add_heading("5. Согласия и локальные акты", level=1)
    _add_paragraph(doc, "До запуска портала организация должна подготовить:", bold=True)
    _add_bullets(
        doc,
        [
            "Политику обработки персональных данных.",
            "ЛНА: порядок сбора документов в электронной форме при приеме на работу.",
            "Бланк согласия на обработку ПДн (отдельный документ).",
            "При необходимости - согласие на передачу ПДн конкретным третьим лицам.",
            "Приказ о назначении ответственного за организацию обработки ПДн.",
            "Уведомление в Роскомнадзор (если нет исключения по ч. 2 ст. 22 152-ФЗ).",
            "Акт определения уровня защищенности ИСПДн.",
            "Договоры поручения с хостингом, КЭДО и иными обработчиками.",
        ],
    )
    _add_paragraph(doc, "Обязательные реквизиты согласия (ч. 4 ст. 9 152-ФЗ):", bold=True)
    _add_bullets(
        doc,
        [
            "ФИО и паспортные данные субъекта ПДн.",
            "Наименование оператора (работодателя).",
            "Цель обработки (трудоустройство / рассмотрение кандидатуры).",
            "Перечень обрабатываемых ПДн и действий с ними.",
            "Срок действия согласия и способ отзыва.",
            "Подпись субъекта ПДн (в т.ч. электронная).",
        ],
    )
    _add_paragraph(
        doc,
        "С 1 сентября 2025 года согласие нельзя включать в анкету, заявление о приеме "
        "или трудовой договор. Это должен быть самостоятельный документ.",
        italic=True,
    )

    doc.add_heading("6. Роли и разграничение доступа", level=1)
    _add_table(
        doc,
        ["Роль", "Права", "Ограничения"],
        [
            ["Соискатель", "Загрузка своих документов по ссылке", "Нет доступа к чужим пакетам"],
            ["HR / кадры", "Просмотр пакетов, смена статуса", "Без права публичной пересылки ссылки"],
            ["Бухгалтерия", "Минимум полей для оформления", "Только по роли и согласию"],
            ["СБ", "Проверка документов", "Только внутри системы"],
            ["Администратор ИС", "Техническое сопровождение", "Без просмотра ПДн без основания"],
        ],
    )
    _add_paragraph(
        doc,
        "Запрещенная практика: пересылка кандидату или коллегам публичной ссылки на папку "
        "с паспортными данными. Правильная практика: внутренний workflow с ролями и журналом доступа.",
        bold=True,
    )

    doc.add_heading("7. Хранение и уничтожение данных", level=1)
    _add_paragraph(
        doc,
        "Используются два контура хранения. Временный контур портала содержит полные сканы "
        "и результаты OCR до принятия решения. Постоянный контур - личное дело и КЭДО после приема.",
    )
    _add_bullets(
        doc,
        [
            "При отказе в приеме: уничтожение ПДн соискателя в течение 30 дней (ч. 4 ст. 21 152-ФЗ).",
            "При отзыве согласия: прекращение обработки и уничтожение в течение 30 дней.",
            "При приеме: перенос в личное дело, архивное хранение 50 лет (дела после 2003 г.).",
            "Каждое уничтожение оформляется актом.",
        ],
    )

    doc.add_heading("8. Проверка подлинности документов", level=1)
    _add_table(
        doc,
        ["Уровень", "Метод", "Правовый статус"],
        [
            ["1", "OCR полей паспорта", "Вспомогательная обработка в рамках согласия"],
            ["2", "Проверка формата, MRZ, шаблона", "Отсев явных подделок"],
            ["3", "Сверка через СМЭВ / гос. API", "Проверка действительности документа"],
            ["4", "Face match / liveness", "Биометрия - отдельное согласие ст. 11"],
        ],
    )
    _add_paragraph(
        doc,
        "Для типового трудоустройства достаточно уровней 1-2 и ручной проверки кадрами. "
        "Уровень 4 применяется только при отдельном правовом и организационном оформлении.",
    )

    doc.add_heading("9. Риски и антипаттерны", level=1)
    _add_table(
        doc,
        ["Нельзя", "Почему", "Как правильно"],
        [
            ["Согласие внутри анкеты", "Ст. 9 152-ФЗ с 01.09.2025", "Отдельный документ согласия"],
            ["Ссылка на пакет в мессенджере", "Утечка, нет контроля", "Внутренний доступ по ролям"],
            ["Запрос лишних документов", "Ст. 65 ТК РФ", "Только законный перечень"],
            ["Хранение вечно в портале", "Нет цели обработки", "Перенос или уничтожение"],
            ["Face match без согласия", "Ст. 11 152-ФЗ", "Отдельное согласие на биометрию"],
            ["Сервер за пределами РФ", "Ст. 18.5 152-ФЗ", "Локализация в РФ"],
        ],
    )
    _add_paragraph(doc, "Штрафы за утечки (420-ФЗ, с 30.05.2025):", bold=True)
    _add_bullets(
        doc,
        [
            "1 000 - 10 000 субъектов: 3-5 млн руб. для юрлица.",
            "10 000 - 100 000 субъектов: 5-10 млн руб.",
            "Свыше 100 000 или утечка биометрии: 15-20 млн руб.",
            "Неуведомление Роскомнадзора об утечке: 1-3 млн руб.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("Приложение A. Макеты экранов", level=1)
    _add_paragraph(
        doc,
        "Ниже представлены wireframe-макеты будущего портала в стиле HUB-IT. "
        "Это проектные экраны, не текущий production UI.",
    )
    _add_image(doc, "mockup-01-landing.png", "Рис. A.1. Вход по персональной ссылке - политика ПДн")
    _add_image(doc, "mockup-02-consent.png", "Рис. A.2. Отдельное согласие на обработку ПДн")
    _add_image(doc, "mockup-03-upload.png", "Рис. A.3. Загрузка документов по ст. 65 ТК РФ")
    _add_image(doc, "mockup-04-hr-dashboard.png", "Рис. A.4. Кабинет кадров - внутренний доступ")

    doc.add_page_break()
    doc.add_heading("Приложение B. Интеграция в HUB-IT", level=1)
    _add_paragraph(
        doc,
        "В текущей кодовой базе HUB-IT отдельного HR-модуля нет. Ниже - как новый портал "
        "может опереться на существующие компоненты платформы.",
    )
    _add_table(
        doc,
        ["Компонент HUB-IT", "Путь в репозитории", "Применение"],
        [
            ["Шифрование ПДн", "WEB-itinvent/backend/services/secret_crypto_service.py", "Хранение паспорта и адреса"],
            ["Модель сотрудника", "WEB-itinvent/backend/appdb/tickets_models.py", "Образец схемы PostgreSQL"],
            ["Права на ПДн", "WEB-itinvent/backend/services/authorization_service.py", "hr.personal_data.read"],
            ["Публичная ссылка", "WEB-itinvent/frontend/src/pages/SharedFile.jsx", "Паттерн токена для кандидата"],
            ["Загрузка файлов", "WEB-itinvent/backend/services/my_files_service.py", "Antivirus, quota, SHA-256"],
            ["Workflow", "WEB-itinvent/backend/services/hub_service.py", "Статусы проверки пакета"],
            ["Оргструктура", "WEB-itinvent/backend/services/department_service.py", "Привязка к отделу"],
        ],
    )
    _add_image(doc, "diagram-05-hub-integration.png", "Рис. B.1. Архитектура интеграции в HUB-IT")

    _add_paragraph(doc, "Планируемые новые элементы (roadmap):", bold=True)
    _add_bullets(
        doc,
        [
            "Backend: api/v1/hr_onboarding.py, services/hr_onboarding_service.py, Alembic-миграция.",
            "Frontend: /onboarding/:token (публичный контур), /hr/onboarding (кадры).",
            "Права: hr.read, hr.manage, hr.personal_data.read.",
            "Retention job: автоудаление пакетов отклоненных кандидатов через 30 дней.",
        ],
    )

    _add_footer(doc)
    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Written: {OUTPUT}")


if __name__ == "__main__":
    main()
