#!/usr/bin/env python3
"""Короткий отчёт о конкурсном задании HUB-IT (DOCX)."""
from __future__ import annotations

import json
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
METRICS_PATH = ROOT / "documentation" / "user-guides" / "_hub_platform_metrics_snapshot.json"
OUTPUT = ROOT / "documentation" / "user-guides" / "hub-platform-competition-report-kozlovsky.docx"

AUTHOR_FIO = "Козловский Максим Евгеньевич"
COMPETITION_TITLE = (
    "Конкурсное задание (видео-визитка): разработка функционального сервиса "
    "для работы, коммуникации и оптимизации бизнес-процессов IT-отдела"
)
REPORT_SUBTITLE = "Единая платформа HUB-IT"

PRIMARY = RGBColor(0x19, 0x76, 0xD2)
MUTED = RGBColor(0x64, 0x74, 0x8B)


def _style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(11)
    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Segoe UI"
        style.font.color.rgb = PRIMARY
        style.font.bold = True


def _add_paragraph(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _fmt_metric(value: Any, fallback: str = "н/д") -> str:
    if isinstance(value, int):
        return f"{value:,}".replace(",", " ")
    if value is None:
        return fallback
    if isinstance(value, dict) and "error" in value:
        return fallback
    return str(value)


def _load_metrics() -> dict[str, Any]:
    if METRICS_PATH.exists():
        return json.loads(METRICS_PATH.read_text(encoding="utf-8"))
    return {"generated_at": date.today().isoformat(), "postgres": {}, "scan_center": {}}


def build_document(metrics: dict[str, Any]) -> Document:
    pg = metrics.get("postgres") or {}
    scan = metrics.get("scan_center") or {}
    generated = str(metrics.get("generated_at", ""))[:10]
    report_date = date.today().strftime("%d.%m.%Y")

    doc = Document()
    _style_document(doc)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Отчёт о выполнении конкурсного задания")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = PRIMARY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run(REPORT_SUBTITLE)
    sub.font.size = Pt(14)
    sub.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Автор: {AUTHOR_FIO}\n"
        f"Дата формирования: {report_date}\n"
        f"Период показателей: с 01.01.2025 по {generated or report_date}"
    )
    meta_run.font.size = Pt(11)

    doc.add_paragraph("")

    doc.add_heading("Сведения об исполнителе", level=1)
    _add_paragraph(doc, f"Фамилия, имя, отчество: {AUTHOR_FIO}")
    _add_paragraph(doc, f"Наименование задания: {COMPETITION_TITLE}")

    doc.add_heading("1. Цель, которую я ставил перед собой", level=1)
    _add_paragraph(
        doc,
        "В видео-визитке и в рамках конкурсного задания я поставил цель создать единый "
        "функциональный сервис для IT-отдела и сотрудников компании: ускорить коммуникации, "
        "организовать постановку и контроль задач, сократить потери времени на согласования "
        "и сделать прозрачным учёт техники и выполненных работ.",
    )
    _add_paragraph(
        doc,
        "До внедрения HUB-IT процессы были разрозненными: переписка шла в личных мессенджерах, "
        "поручения фиксировались устно или в таблицах, акты перемещения техники оформлялись "
        "на бумаге, а данные о компьютерах собирались вручную. Целью было объединить эти "
        "процессы в одной платформе, доступной в браузере, через Telegram-бот и с автоматическим "
        "сбором данных с рабочих ПК.",
    )

    doc.add_heading("2. Действия, предпринятые для достижения результата", level=1)
    _add_paragraph(
        doc,
        "Для реализации цели я спроектировал и внедрил платформу HUB-IT как единый контур "
        "для коммуникаций, управления работой и IT-учёта. Основные направления работы:",
    )
    _add_bullets(
        doc,
        [
            "Коммуникации: корпоративный чат, почта в браузере, push-уведомления, "
            "обсуждение задач в чате, AI-помощники для типовых вопросов.",
            "Постановка и контроль задач: проекты, сроки, приоритеты, чек-листы, "
            "проверка результата руководителем, email-напоминания, аналитика и выгрузка в Excel.",
            "Учёт техники: web и Telegram-бот для поиска, перемещений, актов, QR-кодов, "
            "регистрации обслуживания; распознавание номеров с фото и разбор PDF актов.",
            "Инфраструктура IT: автоматический inventory-агент на ПК, мониторинг компьютеров "
            "и МФУ, карты сетей, справочник ВКС-терминалов.",
            "Безопасность и контроль: двухфакторный вход и passkey, хранилище паролей, "
            "Scan Center для проверки чувствительных документов на рабочих местах.",
            "Качество и эксплуатация: единая база данных, миграции схемы, автотесты, "
            "развёртывание через PM2, адаптивный интерфейс для работы с телефона.",
        ],
    )

    doc.add_heading("3. Достигнутые показатели", level=1)
    _add_paragraph(
        doc,
        "Ниже приведены фактические показатели использования платформы на дату выгрузки "
        f"({generated or report_date}). Цифры получены из рабочей базы данных HUB-IT.",
    )

    metrics_rows = [
        ("Учётные записи в web-приложении", _fmt_metric(pg.get("users_total"))),
        ("Пользователи с привязанным Telegram", _fmt_metric(pg.get("users_with_telegram"))),
        ("Активные пользователи с 2025 года", _fmt_metric(pg.get("active_users_since_2025"))),
        ("Подразделения в структуре компании", _fmt_metric(pg.get("departments_total"))),
        ("Задачи Hub (всего)", _fmt_metric(pg.get("hub_tasks_total"))),
        ("Задачи Hub (выполнено)", _fmt_metric(pg.get("hub_tasks_done"))),
        ("Задачи Hub (создано с 2025 года)", _fmt_metric(pg.get("hub_tasks_since_2025"))),
        ("Комментарии к задачам", _fmt_metric(pg.get("hub_task_comments"))),
        ("Проекты задач", _fmt_metric(pg.get("hub_task_projects"))),
        ("Объявления компании", _fmt_metric(pg.get("hub_announcements"))),
        ("Файлы в «Мои файлы»", _fmt_metric(pg.get("my_files_total"))),
        ("Записи в хранилище паролей", _fmt_metric(pg.get("password_vault_entries"))),
        ("Учтённые компьютеры (inventory)", _fmt_metric(pg.get("inventory_hosts_total"))),
        ("Компьютеры (активны за 7 дней)", _fmt_metric(pg.get("inventory_hosts_recent"))),
        ("Перемещения техники", _fmt_metric(pg.get("equipment_transfers"))),
        ("Регистрации чистки ПК", _fmt_metric(pg.get("pc_cleanings"))),
        ("Замены батарей ИБП", _fmt_metric(pg.get("battery_replacements"))),
        ("Замены комплектующих ПК", _fmt_metric(pg.get("component_replacements"))),
        ("Статьи базы знаний", _fmt_metric(pg.get("kb_articles"))),
        ("Исходящие письма (лог)", _fmt_metric(pg.get("mail_messages_log"))),
    ]

    if pg.get("chat_messages") is not None:
        metrics_rows.extend([
            ("Диалоги в корпоративном чате", _fmt_metric(pg.get("chat_conversations"))),
            ("Сообщения в чате (всего)", _fmt_metric(pg.get("chat_messages"))),
            ("Сообщения в чате (с 2025 года)", _fmt_metric(pg.get("chat_messages_since_2025"))),
        ])

    if scan.get("configured") and isinstance(scan.get("scan_agents"), int):
        metrics_rows.extend([
            ("Scan Center: агенты", _fmt_metric(scan.get("scan_agents"))),
            ("Scan Center: инциденты", _fmt_metric(scan.get("scan_incidents"))),
            ("Scan Center: задания сканирования", _fmt_metric(scan.get("scan_jobs"))),
        ])

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Показатель"
    hdr[1].text = "Значение"
    for label, value in metrics_rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    doc.add_heading("4. Польза для компании", level=1)
    _add_paragraph(
        doc,
        "Результат работы принёс компании измеримую и организационную пользу. "
        "Платформа сократила зависимость от разрозненных каналов связи и бумажных процессов.",
    )
    _add_bullets(
        doc,
        [
            f"Коммуникации: корпоративный чат и единое окно с почтой "
            f"({_fmt_metric(pg.get('mail_messages_log'))} исходящих писем в логе) снижают время "
            "на согласования и поиск переписки по задаче.",
            f"Контроль работы: {_fmt_metric(pg.get('hub_tasks_total'))} поставленных задач, "
            f"из них {_fmt_metric(pg.get('hub_tasks_done'))} закрытых - поручения не теряются, "
            "сроки и ответственные видны руководителю.",
            f"IT-учёт: {_fmt_metric(pg.get('equipment_transfers'))} перемещений техники, "
            f"{_fmt_metric(pg.get('inventory_hosts_total'))} компьютеров в автоматическом учёте - "
            "меньше ручного ввода и быстрее поиск оборудования в поле через бот.",
            f"Сервисные работы: зафиксированы {_fmt_metric(pg.get('pc_cleanings'))} чисток ПК, "
            f"{_fmt_metric(pg.get('battery_replacements'))} замен батарей, "
            f"{_fmt_metric(pg.get('component_replacements'))} замен комплектующих - "
            "прозрачная статистика для планирования и отчётности.",
            "Безопасность: централизованное хранилище паролей, 2FA/passkey и Scan Center "
            "укрепляют контроль доступа и снижают риск утечки чувствительных документов.",
            f"Охват: {_fmt_metric(pg.get('users_total'))} пользователей web, "
            f"{_fmt_metric(pg.get('users_with_telegram'))} пользователей Telegram-бота - "
            "сервис используется и в офисе, и у выездных специалистов.",
            f"Информационная безопасность: Scan Center обработал "
            f"{_fmt_metric(scan.get('scan_jobs'))} заданий сканирования на "
            f"{_fmt_metric(scan.get('scan_agents'))} рабочих ПК.",
        ],
    )

    doc.add_heading("5. Созданный функционал (краткий обзор)", level=1)
    _add_paragraph(
        doc,
        "В рамках конкурсного задания реализована не отдельная программа, а целая платформа "
        "с модулями для разных ролей. Ключевые компоненты:",
    )
    _add_bullets(
        doc,
        [
            "Главная - сводка срочных задач, почты, чата и объявлений.",
            "Задачи - постановка, контроль сроков, проверка, файлы, чек-листы, Гант и аналитика.",
            "Чат и почта - корпоративные коммуникации в одном приложении.",
            "База учёта техники - поиск, перемещения, акты, QR, обслуживание.",
            "Telegram-бот - поиск и передача техники, учёт работ в поле.",
            "Компьютеры, МФУ, сети, ВКС - мониторинг IT-инфраструктуры.",
            "Scan Center, пароли, база знаний, адресная книга, статистика, администрирование.",
        ],
    )

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{AUTHOR_FIO}\n{report_date}")
    run.italic = True

    return doc


def main() -> None:
    metrics = _load_metrics()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document(metrics)
    doc.save(OUTPUT)
    print(f"Written: {OUTPUT}")


if __name__ == "__main__":
    main()
