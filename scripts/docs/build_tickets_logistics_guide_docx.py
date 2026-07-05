#!/usr/bin/env python3
"""Build Tickets/Logistics management guide DOCX for HUB-IT."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
ASSETS = ROOT / "documentation" / "user-guides" / "assets" / "tickets"
OUTPUT = ROOT / "documentation" / "user-guides" / "tickets-logistics-guide.docx"

PRIMARY = RGBColor(0x19, 0x76, 0xD2)
MUTED = RGBColor(0x64, 0x74, 0x8B)
SUCCESS = RGBColor(0x2E, 0x7D, 0x32)


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shading)


def _style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(11)
    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Segoe UI"
        style.font.color.rgb = PRIMARY
        style.font.bold = True
        if level == 1:
            style.font.size = Pt(18)
        elif level == 2:
            style.font.size = Pt(14)
        else:
            style.font.size = Pt(12)


def _add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("HUB-IT | Билеты / Логистика | Конфиденциально")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def _add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(hdr_cells[i], "1976d2")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph("")


def _add_image(doc: Document, filename: str, caption: str, width_in: float = 6.3) -> None:
    path = ASSETS / filename
    if not path.exists():
        _add_paragraph(doc, f"[Диаграмма не найдена: {filename}. Запустите scripts/docs/render_tickets_assets.py]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = MUTED
    doc.add_paragraph("")


STATUS_ROWS = [
    ["Новая", "Заявка поступила, ещё не обработана"],
    ["Проверка данных", "Оператор сверяет ФИО, даты, маршрут, документы"],
    ["Не хватает данных", "Ждём информацию от координатора или сотрудника"],
    ["Готова к покупке", "Все данные есть, можно покупать билет"],
    ["В работе", "Идёт покупка у перевозчика или агентства"],
    ["Куплен", "Билет оформлен, можно прикрепить PDF"],
    ["Нужен обмен", "Требуется переоформление билета"],
    ["Возврат", "Оформляется возврат средств"],
    ["Отмена", "Поездка отменена"],
    ["Не явился", "Сотрудник не воспользовался билетом"],
    ["Закрыта", "Случай завершён"],
    ["Архив", "Убрано из активной работы"],
]

PERMISSION_ROWS = [
    ["Билеты: просмотр", "Доступ к разделу, списку, дашборду, отчётам"],
    ["Билеты: создание и изменения", "Создание заявок, статусы, вложения, импорт, финансы"],
    ["Билеты: персональные данные", "Просмотр паспорта и даты рождения (иначе скрыты)"],
]

SLA_ROWS = [
    ["Вылет скоро", "Вылет через N дней, билет ещё не куплен", "3 дня"],
    ["Не хватает данных", "Заявка долго в статусе «не хватает данных»", "3 дня"],
    ["Нет движения", "Статус не менялся слишком долго", "5 дней"],
    ["Новая потеря", "Зафиксирована финансовая потеря за сутки", "24 часа"],
]

IMPORT_COLOR_ROWS = [
    ["Белый", "Новая"],
    ["Зелёный", "Куплен"],
    ["Жёлтый", "В работе"],
    ["Красный", "Отмена"],
    ["Оранжевый", "Нужен обмен"],
    ["Голубой", "Проверка данных"],
    ["Серый", "Закрыта"],
]


def build_document() -> Document:
    doc = Document()
    _style_document(doc)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Билеты / Логистика")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = PRIMARY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Руководство для руководства и операционных сотрудников")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Внутренний документ | Версия 1.0 | {date.today().strftime('%d.%m.%Y')}\n"
        "Аудитория: руководство, координаторы, операторы логистики"
    )
    meta_run.font.size = Pt(11)

    doc.add_page_break()

    doc.add_heading("Оглавление", level=1)
    toc = [
        "1. Назначение модуля",
        "2. Кто пользователи",
        "3. Права доступа",
        "4. Основные понятия",
        "5. Бизнес-процесс: путь заявки",
        "6. Документооборот",
        "7. Разделы экрана «Билеты»",
        "8. Пошаговые сценарии работы",
        "9. Канбан и дашборд",
        "10. Импорт из Excel",
        "11. Контроль SLA",
        "12. Отчёты",
        "13. Типовой рабочий день",
        "14. Ограничения модуля",
        "15. Краткая памятка",
    ]
    for item in toc:
        doc.add_paragraph(item)

    doc.add_page_break()

    doc.add_heading("1. Назначение модуля", level=1)
    p = doc.add_paragraph()
    status = p.add_run("СТАТУС: ГОТОВ К ЭКСПЛУАТАЦИИ. ")
    status.bold = True
    status.font.color.rgb = SUCCESS
    p.add_run(
        "Модуль ведёт учёт заявок на проезд сотрудников (авиа, ж/д, автобус) "
        "по объектам компании: филиалы, площадки, регионы."
    )
    _add_paragraph(doc, "Система позволяет:", bold=True)
    _add_bullets(
        doc,
        [
            "Фиксировать, кто, куда и когда должен выехать.",
            "Вести заявку от поступления до покупки билета и закрытия.",
            "Хранить документы: маршрут, PDF билета, чеки, ваучеры.",
            "Учитывать потери, возвраты и обмены.",
            "Импортировать накопленные Excel-реестры.",
            "Контролировать сроки через дашборд, канбан и SLA-уведомления.",
        ],
    )
    _add_paragraph(doc, "Важно:", bold=True)
    _add_bullets(
        doc,
        [
            "Модуль работает только в веб-интерфейсе HUB-IT (меню «Билеты»).",
            "Не связан с разделом «Задачи» (Hub Tasks) и с Telegram-ботом.",
            "Покупка билета у перевозчика выполняется вне системы; здесь ведётся учёт и контроль.",
        ],
    )

    doc.add_heading("2. Кто пользователи", level=1)
    _add_table(
        doc,
        ["Роль", "Ответственность", "Типичные задачи"],
        [
            ["Руководитель / контролёр", "Контроль сроков и рисков", "Дашборд, отчёты, SLA"],
            ["Оператор логистики", "Обработка заявок", "Статусы, вложения, покупка, финансы"],
            ["Координатор объекта", "Данные по поездкам", "Создание заявок, срочность, ФИО"],
            ["Администратор системы", "Настройка", "Объекты, права, правила SLA"],
        ],
    )
    _add_image(doc, "diagram-01-roles.png", "Рис. 1. Роли и зоны ответственности")

    doc.add_heading("3. Права доступа", level=1)
    _add_paragraph(
        doc,
        "Права выдаёт администратор в настройках учётной записи. Без права «просмотр» "
        "раздел «Билеты» в меню не отображается.",
    )
    _add_table(doc, ["Право", "Что даёт"], PERMISSION_ROWS)

    doc.add_heading("4. Основные понятия", level=1)
    _add_table(
        doc,
        ["Термин", "Значение"],
        [
            ["Объект", "Филиал, площадка или регион, к которому относится заявка"],
            ["Сотрудник", "Человек, для которого покупается билет (справочник модуля)"],
            ["Заявка", "Главный документ: одна поездка или запрос на билет"],
            ["Статус", "Этап обработки заявки"],
            ["Исполнитель", "Пользователь системы, который ведёт заявку"],
            ["Вложение", "Файл: маршрут, PDF билета, чек, ваучер"],
            ["Финансовая операция", "Потеря, возврат или обмен денег по билету"],
            ["SLA-уведомление", "Автоматическое предупреждение о просрочке или риске"],
        ],
    )

    doc.add_page_break()

    doc.add_heading("5. Бизнес-процесс: путь заявки", level=1)
    _add_paragraph(
        doc,
        "Ниже показан типовой путь заявки от поступления до архива. "
        "Оператор переводит заявку только по разрешённым переходам. "
        "Администратор может обходить ограничения при необходимости.",
    )
    _add_image(doc, "diagram-02-lifecycle.png", "Рис. 2. Жизненный цикл заявки на билет")
    _add_table(doc, ["Статус на экране", "Смысл для бизнеса"], STATUS_ROWS)

    doc.add_heading("6. Документооборот", level=1)
    _add_image(doc, "diagram-03-documents.png", "Рис. 3. Документооборот и данные")
    _add_paragraph(doc, "Типы вложений к заявке:", bold=True)
    _add_table(
        doc,
        ["Тип", "Когда прикрепляют"],
        [
            ["Маршрут", "После согласования маршрута"],
            ["PDF билета", "После покупки"],
            ["Чек / квитанция", "После оплаты"],
            ["Ваучер", "При бронировании через агентство"],
            ["Прочее", "Любые дополнительные материалы"],
        ],
    )
    _add_paragraph(
        doc,
        "Ограничения: до 10 файлов на заявку, размер одного файла до 20 МБ. "
        "Форматы: PDF, JPG, PNG, DOC, DOCX, XLS, XLSX.",
    )
    _add_paragraph(
        doc,
        "Паспортные данные и дата рождения хранятся в зашифрованном виде. "
        "Просмотр доступен только пользователям с отдельным правом.",
        italic=True,
    )

    doc.add_page_break()

    doc.add_heading("7. Разделы экрана «Билеты»", level=1)
    _add_paragraph(doc, "После входа в раздел доступны 7 вкладок:", bold=True)
    _add_table(
        doc,
        ["Вкладка", "Для кого", "Назначение"],
        [
            ["Список", "Все с правом просмотра", "Таблица заявок и карточка выбранной заявки"],
            ["Канбан", "Операторы, руководство", "Наглядная доска по этапам"],
            ["Дашборд", "Руководство", "Сводные цифры и показатели"],
            ["Отчёты", "Руководство, бухгалтерия", "Потери, возвраты, обмены"],
            ["Импорт", "Операторы", "Загрузка Excel-реестров"],
            ["Справочники", "Администратор, координаторы", "Объекты и сотрудники"],
            ["Правила SLA", "Администратор", "Настройка контрольных сроков"],
        ],
    )
    _add_paragraph(
        doc,
        "В верхней части страницы отображаются активные SLA-уведомления (если есть).",
    )

    doc.add_heading("8. Пошаговые сценарии работы", level=1)

    doc.add_heading("8.1. Координатор: создать заявку", level=2)
    _add_numbered(
        doc,
        [
            "Открыть «Билеты» и нажать «Создать заявку».",
            "Выбрать сотрудника из справочника или ввести ФИО нового.",
            "Выбрать объект (филиал).",
            "Указать даты вылета и прибытия, маршрут, стоимость (если известна).",
            "При срочности включить «Срочная заявка».",
            "Нажать «Создать». Заявка появится в статусе «Новая».",
        ],
    )

    doc.add_heading("8.2. Оператор: обработать заявку", level=2)
    _add_numbered(
        doc,
        [
            "Открыть вкладку «Список» или «Канбан».",
            "Выбрать заявку и открыть карточку.",
            "Проверить данные сотрудника и маршрут.",
            "Нажать «Сменить статус» и выбрать следующий этап.",
            "При необходимости добавить комментарий.",
            "После покупки прикрепить PDF билета и чек.",
            "Перевести в «Куплен», затем в «Закрыта» после поездки.",
        ],
    )

    doc.add_heading("8.3. Если не хватает данных", level=2)
    _add_numbered(
        doc,
        [
            "Перевести заявку в «Не хватает данных».",
            "Оставить комментарий с перечнем недостающего.",
            "Координатор дополняет данные в справочнике или заявке.",
            "После получения данных вернуть в «Проверка данных».",
        ],
    )

    doc.add_heading("8.4. Возврат, обмен или потеря", level=2)
    _add_bullets(
        doc,
        [
            "Через статус: «Нужен обмен» -> «В работе» -> «Куплен»; «Возврат» -> «Закрыта».",
            "Через финансовые операции (вкладка «Отчёты»): зафиксировать потерю, возврат или обмен с суммой и причиной.",
            "Новая потеря за сутки попадает в SLA-уведомления.",
        ],
    )

    doc.add_heading("8.5. Руководитель: ежедневный контроль", level=2)
    _add_numbered(
        doc,
        [
            "Открыть «Дашборд»: активные заявки, вылеты сегодня/завтра/3 дня, проблемные объекты.",
            "Просмотреть SLA-уведомления в шапке страницы.",
            "При необходимости открыть «Отчёты» и выгрузить потери за период в Excel.",
        ],
    )

    doc.add_heading("8.6. Администратор: объекты", level=2)
    _add_numbered(
        doc,
        [
            "Вкладка «Справочники» -> блок объектов.",
            "Создать объект: код, название, регион.",
            "При необходимости отключить неактуальный объект.",
        ],
    )
    _add_paragraph(doc, "Создание объектов доступно только администратору.", italic=True)

    doc.add_page_break()

    doc.add_heading("9. Канбан и дашборд", level=1)
    _add_image(doc, "diagram-05-kanban.png", "Рис. 5. Канбан-доска: колонки этапов")
    _add_paragraph(doc, "Колонки канбана:", bold=True)
    _add_bullets(
        doc,
        [
            "Не запущен - новая, проверка, не хватает данных, готова к покупке.",
            "В работе - идёт покупка билета.",
            "Куплен - билет оформлен.",
            "Возврат / обмен - переоформление или возврат.",
            "Отмена - отменённые заявки.",
            "Проблема - не явился, срочные, требует проверки.",
        ],
    )
    _add_paragraph(doc, "Дашборд показывает:", bold=True)
    _add_bullets(
        doc,
        [
            "Число активных заявок и разбивку по статусам.",
            "Вылеты сегодня, завтра и в ближайшие 3 дня.",
            "Топ объектов с проблемами и топ исполнителей по нагрузке.",
        ],
    )

    doc.add_heading("10. Импорт из Excel", level=1)
    _add_paragraph(
        doc,
        "Импорт предназначен для переноса накопленных таблиц при запуске или миграции с Excel-учёта. "
        "Для ежедневной работы рекомендуется создавать заявки вручную.",
    )
    _add_image(doc, "diagram-04-import.png", "Рис. 4. Импорт заявок из Excel")
    _add_paragraph(doc, "Шаги для оператора:", bold=True)
    _add_numbered(
        doc,
        [
            "Вкладка «Импорт» -> «Загрузить .xlsx» (до 50 МБ).",
            "Проверить предпросмотр листов и строк.",
            "Для каждого листа указать объект (филиал).",
            "При необходимости скорректировать соответствие цветов и статусов.",
            "Выбрать стратегию дубликатов (обычно «пропустить»).",
            "Нажать «Запустить импорт».",
        ],
    )
    _add_paragraph(doc, "Типовые цвета ячеек при импорте:", bold=True)
    _add_table(doc, ["Цвет в Excel", "Статус в системе"], IMPORT_COLOR_ROWS)

    doc.add_heading("11. Контроль SLA", level=1)
    _add_image(doc, "diagram-06-sla.png", "Рис. 6. Контроль SLA и уведомления")
    _add_table(doc, ["Правило", "Смысл", "Порог по умолчанию"], SLA_ROWS)
    _add_paragraph(
        doc,
        "Пороги настраиваются администратором на вкладке «Правила SLA». "
        "Уведомления показываются на странице «Билеты»; после обработки их можно скрыть.",
    )

    doc.add_heading("12. Отчёты", level=1)
    _add_paragraph(doc, "Отчёт по потерям:", bold=True)
    _add_bullets(
        doc,
        [
            "Фильтр по периоду, объекту и типу операции.",
            "Итоговые суммы по потерям, возвратам и обменам.",
            "Выгрузка в Excel для бухгалтерии.",
        ],
    )
    _add_paragraph(doc, "Экспорт заявок:", bold=True)
    _add_paragraph(doc, "Из списка заявок можно выгрузить отфильтрованный реестр в Excel (нужно право на изменения).")

    doc.add_heading("13. Типовой рабочий день", level=1)
    _add_table(
        doc,
        ["Время", "Координатор", "Оператор", "Руководитель"],
        [
            ["Утро", "Создаёт заявки, помечает срочные", "Проверяет SLA, берёт «новые»", "Смотрит дашборд и вылеты"],
            ["День", "Дополняет данные по запросу", "Покупает билеты, меняет статусы", "Контролирует «проблемную» колонку"],
            ["Вечер", "-", "Закрывает завершённые поездки", "При необходимости - отчёт по потерям"],
        ],
    )

    doc.add_heading("14. Ограничения модуля", level=1)
    _add_table(
        doc,
        ["Ожидание", "Реальность"],
        [
            ["Покупка билета в системе", "Нет - покупка у перевозчика вне системы"],
            ["Интеграция с Telegram-ботом", "Нет - только веб-интерфейс"],
            ["Связь с разделом «Задачи»", "Нет - отдельный контур"],
            ["Email/SMS-рассылка уведомлений", "Нет - уведомления внутри раздела «Билеты»"],
        ],
    )

    doc.add_heading("15. Краткая памятка", level=1)
    _add_numbered(
        doc,
        [
            "Заявка - центральный объект; всё ведётся вокруг неё.",
            "Статус отражает этап; в основном меняет оператор.",
            "Документы прикрепляются к заявке после покупки.",
            "Канбан и дашборд - для руководства и оперативного контроля.",
            "SLA - страховка от «забытых» заявок.",
            "Импорт - для перехода с Excel, не для ежедневной рутины.",
        ],
    )

    _add_footer(doc)
    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Written: {OUTPUT}")


if __name__ == "__main__":
    main()
