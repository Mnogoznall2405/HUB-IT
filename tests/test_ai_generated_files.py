from __future__ import annotations

import json
import sys
from pathlib import Path


PROJECT_ROOT = Path(__file__).resolve().parents[1]
WEB_ROOT = PROJECT_ROOT / "WEB-itinvent"

if str(WEB_ROOT) not in sys.path:
    sys.path.insert(0, str(WEB_ROOT))


def _make_context(*, allow_generated_artifacts: bool = True):
    from backend.ai_chat.tools.context import AI_TOOL_FILES_CREATE, AI_TOOL_FILES_REPORT, AiToolExecutionContext

    return AiToolExecutionContext(
        bot_id="bot-test",
        bot_title="AI Assistant",
        conversation_id="conv-test",
        run_id="run-test",
        user_id=5,
        user_payload={"id": 5, "role": "viewer", "username": "operator"},
        effective_database_id="ITINVENT",
        enabled_tools=[AI_TOOL_FILES_CREATE, AI_TOOL_FILES_REPORT],
        tool_settings={"multi_db_mode": "single", "allowed_databases": []},
        allow_generated_artifacts=allow_generated_artifacts,
    )


def test_generated_file_builder_supports_v1_formats():
    from backend.ai_chat.artifact_generator import build_generated_uploads

    uploads = build_generated_uploads([
        {"format": "xlsx", "file_name": "report.xlsx", "rows": [["name"], ["Ноутбук"]]},
        {"format": "csv", "file_name": "report.csv", "rows": [["name"], ["Ноутбук"]]},
        {"format": "docx", "file_name": "report.docx", "content": "Документ"},
        {"format": "pdf", "file_name": "report.pdf", "content": "PDF с кириллицей"},
        {"format": "txt", "file_name": "report.txt", "content": "Текст"},
        {"format": "md", "file_name": "report.md", "content": "# Markdown"},
        {"format": "json", "file_name": "report.json", "content": {"name": "Ноутбук"}},
    ])

    assert [item.filename for item in uploads] == [
        "report.xlsx",
        "report.csv",
        "report.docx",
        "report.pdf",
        "report.txt",
        "report.md",
        "report.json",
    ]
    assert all(item.file.getbuffer().nbytes > 0 for item in uploads)
    json_upload = uploads[-1]
    assert json.loads(json_upload.file.getvalue().decode("utf-8")) == {"name": "Ноутбук"}
    for upload in uploads:
        upload.file.close()


def test_generated_file_builder_rejects_file_and_row_limits():
    import pytest

    from backend.ai_chat.artifact_generator import GeneratedFileError, build_generated_uploads

    with pytest.raises(GeneratedFileError):
        build_generated_uploads([{"format": "txt", "file_name": f"{index}.txt", "content": "x"} for index in range(11)])

    build_generated_uploads([{"format": "csv", "file_name": "now-allowed.csv", "rows": [["x"]] * 5001}])[0].file.close()

    with pytest.raises(GeneratedFileError):
        build_generated_uploads([{"format": "csv", "file_name": "too-many.csv", "rows": [["x"]] * 50001}])


def test_generated_xlsx_sanitizes_invalid_duplicate_and_long_sheet_titles():
    from openpyxl import load_workbook

    from backend.ai_chat.artifact_generator import build_generated_uploads, normalize_generated_file_specs

    specs = normalize_generated_file_specs([
        {
            "format": "xlsx",
            "file_name": "bad-sheets.xlsx",
            "sheets": [
                {"title": "Genracena: 55", "rows": [["name"], ["row 1"]]},
                {"title": "A:B", "rows": [["name"], ["row 2"]]},
                {"title": "A/B", "rows": [["name"], ["row 3"]]},
                {"title": "Very long sheet name that should still be unique A", "rows": [["name"], ["row 4"]]},
                {"title": "Very long sheet name that should still be unique B", "rows": [["name"], ["row 5"]]},
            ],
        }
    ])

    sheets = specs[0]["sheets"]
    assert [item["title"] for item in sheets] == [
        "Genracena_ 55",
        "A_B",
        "A_B_2",
        "Very long sheet name that shoul",
        "Very long sheet name that sho_2",
    ]

    uploads = build_generated_uploads(specs)
    workbook = load_workbook(uploads[0].file)
    assert workbook.sheetnames == [item["title"] for item in sheets]
    uploads[0].file.close()


def test_generated_xlsx_sanitizes_fallback_title():
    from openpyxl import load_workbook

    from backend.ai_chat.artifact_generator import build_generated_uploads

    uploads = build_generated_uploads([
        {
            "format": "xlsx",
            "file_name": "fallback-title.xlsx",
            "title": "Fallback: title / with bad chars",
            "rows": [["name"], ["row 1"]],
        }
    ])
    workbook = load_workbook(uploads[0].file)
    assert workbook.sheetnames == ["Fallback_ title _ with bad char"]
    uploads[0].file.close()


def test_generated_equipment_rows_enforce_serial_column_when_missing():
    from openpyxl import load_workbook

    from backend.ai_chat.artifact_generator import build_generated_uploads, normalize_generated_file_specs

    serial_header = "\u0421\u0435\u0440\u0438\u0439\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440"
    specs = normalize_generated_file_specs([
        {
            "format": "xlsx",
            "file_name": "equipment.xlsx",
            "rows": [
                {
                    "inv_no": "101",
                    "type_name": "Laptop",
                    "model_name": "Dell Latitude",
                    "owner_name": "Ivanov I.I.",
                    "status": "In use",
                    "branch": "HQ",
                    "location": "101",
                }
            ],
        }
    ])

    rows = specs[0]["rows"]
    assert rows[0] == [
        "inv_no",
        serial_header,
        "type_name",
        "model_name",
        "owner_name",
        "status",
        "branch",
        "location",
    ]
    assert rows[1][1] == "\u2014"

    uploads = build_generated_uploads(specs)
    workbook = load_workbook(uploads[0].file)
    sheet = workbook.active
    assert sheet["B1"].value == serial_header
    assert sheet["B2"].value == "\u2014"
    uploads[0].file.close()


def test_generated_equipment_rows_preserve_serial_and_extra_fields():
    from backend.ai_chat.artifact_generator import normalize_generated_file_specs

    specs = normalize_generated_file_specs([
        {
            "format": "csv",
            "file_name": "equipment.csv",
            "rows": [
                {
                    "inv_no": "101",
                    "serial_no": "SN-101",
                    "model_name": "Dell Latitude",
                    "ip_address": "10.0.0.10",
                }
            ],
        }
    ])

    rows = specs[0]["rows"]
    assert rows[1][1] == "SN-101"
    assert rows[0][-1] == "ip_address"
    assert rows[1][-1] == "10.0.0.10"


def test_generated_list_rows_keep_llm_columns_and_add_missing_serial():
    from backend.ai_chat.artifact_generator import normalize_generated_file_specs

    specs = normalize_generated_file_specs([
        {
            "format": "csv",
            "file_name": "equipment.csv",
            "rows": [
                ["\u0418\u043d\u0432", "\u041c\u043e\u0434\u0435\u043b\u044c", "\u041a\u0430\u0431\u0438\u043d\u0435\u0442"],
                ["101", "Dell Latitude", "101"],
            ],
        }
    ])

    assert specs[0]["rows"] == [
        ["\u0418\u043d\u0432", "\u0421\u0435\u0440\u0438\u0439\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440", "\u041c\u043e\u0434\u0435\u043b\u044c", "\u041a\u0430\u0431\u0438\u043d\u0435\u0442"],
        ["101", "\u2014", "Dell Latitude", "101"],
    ]


def test_generated_flat_equipment_rows_are_recovered_to_table():
    from openpyxl import load_workbook

    from backend.ai_chat.artifact_generator import build_generated_uploads, normalize_generated_file_specs

    headers = [
        "\u0418\u043d\u0432. \u043d\u043e\u043c\u0435\u0440",
        "\u0421\u0435\u0440\u0438\u0439\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440",
        "\u0422\u0438\u043f",
        "\u041c\u043e\u0434\u0435\u043b\u044c",
        "\u0421\u043e\u0442\u0440\u0443\u0434\u043d\u0438\u043a",
        "\u0421\u0442\u0430\u0442\u0443\u0441",
        "\u0424\u0438\u043b\u0438\u0430\u043b",
        "\u041b\u043e\u043a\u0430\u0446\u0438\u044f",
    ]
    specs = normalize_generated_file_specs([
        {
            "format": "xlsx",
            "file_name": "equipment.xlsx",
            "rows": headers
            + ["100665.0", "PC25J7G6", "\u0421\u0438\u0441\u0442\u0435\u043c\u043d\u044b\u0439 \u0431\u043b\u043e\u043a", "Lenovo ThinkCentre M720q", "\u041a\u043e\u0437\u043b\u043e\u0432\u0441\u043a\u0438\u0439 \u041c\u0430\u043a\u0441\u0438\u043c", "\u0420\u0430\u0431\u043e\u0442\u0430\u0435\u0442", "\u0422\u044e\u043c\u0435\u043d\u044c", "19_105"]
            + ["101640.0", "MMT4BEE", "\u041c\u043e\u043d\u0438\u0442\u043e\u0440", "Acer R240HY", "\u041a\u043e\u0437\u043b\u043e\u0432\u0441\u043a\u0438\u0439 \u041c\u0430\u043a\u0441\u0438\u043c", "\u0420\u0430\u0431\u043e\u0442\u0430\u0435\u0442", "\u0422\u044e\u043c\u0435\u043d\u044c", "19_105"],
        }
    ])

    assert len(specs[0]["rows"]) == 3
    assert specs[0]["rows"][0] == headers
    assert specs[0]["rows"][1][0] == "100665.0"
    assert specs[0]["rows"][2][1] == "MMT4BEE"

    uploads = build_generated_uploads(specs)
    workbook = load_workbook(uploads[0].file, read_only=True)
    sheet = workbook.active
    assert sheet.max_row == 3
    assert sheet.max_column == 8
    uploads[0].file.close()


def test_generated_unrecoverable_flat_rows_are_rejected():
    import pytest

    from backend.ai_chat.artifact_generator import GeneratedFileError, normalize_generated_file_specs

    with pytest.raises(GeneratedFileError, match="not a flat list"):
        normalize_generated_file_specs([
            {
                "format": "csv",
                "file_name": "bad.csv",
                "rows": ["\u0418\u043d\u0432. \u043d\u043e\u043c\u0435\u0440", "\u041c\u043e\u0434\u0435\u043b\u044c", "101"],
            }
        ])


def test_generated_flat_rows_with_columns_are_chunked_into_table():
    from openpyxl import load_workbook

    from backend.ai_chat.artifact_generator import build_generated_uploads
    from backend.ai_chat.tools.files import FilesReportArgs, FilesReportTool

    result = FilesReportTool().execute(
        context=_make_context(),
        args=FilesReportArgs.model_validate({
            "format": "xlsx",
            "file_name": "mfu.xlsx",
            "title": "\u0421\u043f\u0438\u0441\u043e\u043a \u041c\u0424\u0423",
            "tables": [
                {
                    "title": "\u041c\u0424\u0423",
                    "columns": [
                        "\u0418\u043d\u0432. \u043d\u043e\u043c\u0435\u0440",
                        "\u0421\u0435\u0440\u0438\u0439\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440",
                        "\u0422\u0438\u043f",
                        "\u041c\u043e\u0434\u0435\u043b\u044c",
                        "\u0421\u043e\u0442\u0440\u0443\u0434\u043d\u0438\u043a",
                        "\u0421\u0442\u0430\u0442\u0443\u0441",
                        "\u0424\u0438\u043b\u0438\u0430\u043b",
                        "\u041b\u043e\u043a\u0430\u0446\u0438\u044f",
                    ],
                    "rows": [
                        "100121.0",
                        "CNBKMBP6SJ",
                        "\u041c\u0424\u0423",
                        "HP LaserJet Pro MFP M428w",
                        "\u041f\u044b\u043b\u0430\u0435\u0432\u0430 \u0412\u0438\u043a\u0442\u043e\u0440\u0438\u044f",
                        "\u0420\u0430\u0431\u043e\u0442\u0430\u0435\u0442",
                        "\u0433.\u0422\u044e\u043c\u0435\u043d\u044c, \u041f\u0435\u0440\u0432\u043e\u043c\u0430\u0439\u0441\u043a\u0430\u044f 19",
                        "19_306",
                        "101287.0",
                        "CNBKM8BIT6",
                        "\u041c\u0424\u0423",
                        "HP MFP M428dw",
                        "\u041a\u0438\u0441\u0435\u043b\u0435\u0432 \u041a\u043e\u043d\u0441\u0442\u0430\u043d\u0442\u0438\u043d",
                        "\u0420\u0430\u0431\u043e\u0442\u0430\u0435\u0442",
                        "\u0433.\u0422\u044e\u043c\u0435\u043d\u044c, \u0413\u0435\u0440\u0446\u0435\u043d\u0430 55",
                        "01",
                    ],
                }
            ],
        }),
    ).to_payload()

    assert result["ok"] is True
    specs = result["data"]["files"]
    table_rows = specs[0]["metadata"]["report_tables"][0]["rows"]
    assert len(table_rows) == 3
    assert table_rows[0] == [
        "\u0418\u043d\u0432. \u043d\u043e\u043c\u0435\u0440",
        "\u0421\u0435\u0440\u0438\u0439\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440",
        "\u0422\u0438\u043f",
        "\u041c\u043e\u0434\u0435\u043b\u044c",
        "\u0421\u043e\u0442\u0440\u0443\u0434\u043d\u0438\u043a",
        "\u0421\u0442\u0430\u0442\u0443\u0441",
        "\u0424\u0438\u043b\u0438\u0430\u043b",
        "\u041b\u043e\u043a\u0430\u0446\u0438\u044f",
    ]
    assert table_rows[1][0] == "100121.0"
    assert table_rows[2][1] == "CNBKM8BIT6"

    uploads = build_generated_uploads(specs)
    workbook = load_workbook(uploads[0].file, read_only=True)
    sheet = workbook["\u041c\u0424\u0423"]
    assert sheet.max_row == 3
    assert sheet.max_column == 8
    uploads[0].file.close()


def test_generated_equipment_rows_drop_blank_trailing_rows():
    from backend.ai_chat.artifact_generator import normalize_generated_file_specs

    specs = normalize_generated_file_specs([
        {
            "format": "csv",
            "file_name": "equipment.csv",
            "rows": [
                {"inv_no": "101", "serial_no": "SN-101", "model_name": "Dell Latitude"},
                {},
                {"inv_no": "", "serial_no": "", "model_name": ""},
                [],
            ],
        }
    ])

    rows = specs[0]["rows"]
    assert len(rows) == 2
    assert rows[1][0] == "101"
    assert rows[1][1] == "SN-101"


def test_generated_non_equipment_rows_do_not_gain_serial_column():
    from backend.ai_chat.artifact_generator import normalize_generated_file_specs

    specs = normalize_generated_file_specs([
        {
            "format": "csv",
            "file_name": "tasks.csv",
            "rows": [{"task": "Call client", "status": "open"}],
        }
    ])

    assert specs[0]["rows"] == [["task", "status"], ["Call client", "open"]]


def test_ai_files_create_tool_returns_normalized_file_specs():
    from backend.ai_chat.tools.files import FilesCreateArgs, FilesCreateTool

    result = FilesCreateTool().execute(
        context=_make_context(),
        args=FilesCreateArgs.model_validate({
            "files": [
                {
                    "format": "excel",
                    "file_name": "Отчет",
                    "rows": [{"Инв": "101", "Модель": "Dell"}],
                }
            ]
        }),
    ).to_payload()

    assert result["ok"] is True
    assert result["data"]["count"] == 1
    assert result["data"]["files"][0]["format"] == "xlsx"
    assert result["data"]["files"][0]["file_name"].endswith(".xlsx")
    assert result["data"]["files"][0]["rows"] == [
        [
            "\u0418\u043d\u0432",
            "\u0421\u0435\u0440\u0438\u0439\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440",
            "\u041c\u043e\u0434\u0435\u043b\u044c",
        ],
        ["101", "\u2014", "Dell"],
    ]


def test_runtime_file_specs_use_full_equipment_tool_result_when_llm_file_is_partial():
    service = __import__("backend.ai_chat.service", fromlist=["_extract_generated_file_specs_from_tool_results"])
    from backend.ai_chat.tools.files import FilesReportArgs, FilesReportTool

    live_result = {
        "tool_id": "itinvent.employee.list_equipment",
        "ok": True,
        "data": {
            "items": [
                {
                    "inv_no": "100665.0",
                    "serial_no": "PC25J7G6",
                    "type_name": "\u0421\u0438\u0441\u0442\u0435\u043c\u043d\u044b\u0439 \u0431\u043b\u043e\u043a",
                    "model_name": "Lenovo ThinkCentre M720q",
                    "owner_name": "\u041a\u043e\u0437\u043b\u043e\u0432\u0441\u043a\u0438\u0439 \u041c\u0430\u043a\u0441\u0438\u043c",
                    "status": "\u0420\u0430\u0431\u043e\u0442\u0430\u0435\u0442",
                    "branch": "\u0433.\u0422\u044e\u043c\u0435\u043d\u044c, \u041f\u0435\u0440\u0432\u043e\u043c\u0430\u0439\u0441\u043a\u0430\u044f 19",
                    "location": "19_105",
                },
                {
                    "inv_no": "101640.0",
                    "serial_no": "MMT4BEE00C237047CA2421",
                    "type_name": "\u041c\u043e\u043d\u0438\u0442\u043e\u0440",
                    "model_name": "Acer R240HY",
                    "owner_name": "\u041a\u043e\u0437\u043b\u043e\u0432\u0441\u043a\u0438\u0439 \u041c\u0430\u043a\u0441\u0438\u043c",
                    "status": "\u0420\u0430\u0431\u043e\u0442\u0430\u0435\u0442",
                    "branch": "\u0433.\u0422\u044e\u043c\u0435\u043d\u044c, \u041f\u0435\u0440\u0432\u043e\u043c\u0430\u0439\u0441\u043a\u0430\u044f 19",
                    "location": "19_105",
                },
            ]
        },
    }
    partial_file_result = FilesReportTool().execute(
        context=_make_context(),
        args=FilesReportArgs.model_validate({
            "format": "xlsx",
            "file_name": "Kozlovsky_Maxim_Equipment.xlsx",
            "title": "\u0422\u0435\u0445\u043d\u0438\u043a\u0430 \u041a\u043e\u0437\u043b\u043e\u0432\u0441\u043a\u043e\u0433\u043e",
            "tables": [
                {
                    "title": "\u041e\u0431\u043e\u0440\u0443\u0434\u043e\u0432\u0430\u043d\u0438\u0435",
                    "columns": [
                        "\u0418\u043d\u0432. \u043d\u043e\u043c\u0435\u0440",
                        "\u0421\u0435\u0440\u0438\u0439\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440",
                        "\u0422\u0438\u043f",
                        "\u041c\u043e\u0434\u0435\u043b\u044c",
                        "\u0421\u043e\u0442\u0440\u0443\u0434\u043d\u0438\u043a",
                        "\u0421\u0442\u0430\u0442\u0443\u0441",
                        "\u0424\u0438\u043b\u0438\u0430\u043b",
                        "\u041b\u043e\u043a\u0430\u0446\u0438\u044f",
                    ],
                    "rows": [
                        "101640.0",
                        "MMT4BEE00C237047CA2421",
                        "\u041c\u043e\u043d\u0438\u0442\u043e\u0440",
                        "Acer R240HY",
                        "\u041a\u043e\u0437\u043b\u043e\u0432\u0441\u043a\u0438\u0439 \u041c\u0430\u043a\u0441\u0438\u043c",
                        "\u0420\u0430\u0431\u043e\u0442\u0430\u0435\u0442",
                        "\u0433.\u0422\u044e\u043c\u0435\u043d\u044c, \u041f\u0435\u0440\u0432\u043e\u043c\u0430\u0439\u0441\u043a\u0430\u044f 19",
                        "19_105",
                    ],
                }
            ],
        }),
    ).to_payload()

    specs = service._extract_generated_file_specs_from_tool_results([live_result, partial_file_result])

    assert len(specs) == 1
    rows = specs[0]["metadata"]["report_tables"][0]["rows"]
    assert len(rows) == 3
    assert rows[0] == [
        "\u0418\u043d\u0432. \u043d\u043e\u043c\u0435\u0440",
        "\u0421\u0435\u0440\u0438\u0439\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440",
        "\u0422\u0438\u043f",
        "\u041c\u043e\u0434\u0435\u043b\u044c",
        "\u0421\u043e\u0442\u0440\u0443\u0434\u043d\u0438\u043a",
        "\u0421\u0442\u0430\u0442\u0443\u0441",
        "\u0424\u0438\u043b\u0438\u0430\u043b",
        "\u041b\u043e\u043a\u0430\u0446\u0438\u044f",
    ]
    assert rows[1][0] == "100665.0"
    assert rows[2][0] == "101640.0"
    assert specs[0]["sheets"][-1]["rows"] == rows


def test_ai_files_create_tool_accepts_single_file_and_fallback_name():
    from backend.ai_chat.tools.files import FilesCreateArgs, FilesCreateTool

    args = FilesCreateArgs.model_validate({
        "files": {
            "file_name": "",
            "rows": [["\u0418\u043d\u0432", "\u041c\u043e\u0434\u0435\u043b\u044c"], ["101", "Dell"]],
            "metadata": "ignored",
        }
    })
    result = FilesCreateTool().execute(context=_make_context(), args=args).to_payload()

    assert result["ok"] is True
    spec = result["data"]["files"][0]
    assert spec["format"] == "xlsx"
    assert spec["file_name"] == "generated-file.xlsx"
    assert spec["rows"] == [
        ["\u0418\u043d\u0432", "\u0421\u0435\u0440\u0438\u0439\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440", "\u041c\u043e\u0434\u0435\u043b\u044c"],
        ["101", "\u2014", "Dell"],
    ]


def test_ai_files_create_tool_respects_generated_files_master_switch():
    from backend.ai_chat.tools.files import FilesCreateArgs, FilesCreateTool

    result = FilesCreateTool().execute(
        context=_make_context(allow_generated_artifacts=False),
        args=FilesCreateArgs.model_validate({
            "files": [{"format": "txt", "file_name": "blocked.txt", "content": "blocked"}]
        }),
    ).to_payload()

    assert result["ok"] is False
    assert "disabled" in result["error"]


def test_ai_files_report_tool_builds_polished_xlsx_report():
    from openpyxl import load_workbook

    from backend.ai_chat.artifact_generator import build_generated_uploads
    from backend.ai_chat.tools.files import FilesReportArgs, FilesReportTool

    result = FilesReportTool().execute(
        context=_make_context(),
        args=FilesReportArgs.model_validate({
            "format": "xlsx",
            "file_name": "inventory-report",
            "title": "Отчёт по инвентарю",
            "summary": "Найдено 2 устройства.",
            "sections": [{"heading": "Итог", "body": "Все устройства в эксплуатации."}],
            "tables": [
                {
                    "title": "Ноутбуки",
                    "rows": [
                        {"Инв": "101", "Модель": "Dell Latitude"},
                        {"Инв": "102", "Модель": "HP EliteBook"},
                    ],
                }
            ],
        }),
    ).to_payload()

    assert result["ok"] is True
    spec = result["data"]["files"][0]
    assert spec["file_name"] == "inventory-report.xlsx"
    assert spec["sheets"][0]["title"] == "\u0421\u0432\u043e\u0434\u043a\u0430"
    assert spec["sheets"][1]["title"] == "\u041d\u043e\u0443\u0442\u0431\u0443\u043a\u0438"

    uploads = build_generated_uploads([spec])
    workbook = load_workbook(uploads[0].file)
    assert workbook.sheetnames[0] == "\u0421\u0432\u043e\u0434\u043a\u0430"
    assert workbook.sheetnames[1] == "\u041d\u043e\u0443\u0442\u0431\u0443\u043a\u0438"
    sheet = workbook["\u041d\u043e\u0443\u0442\u0431\u0443\u043a\u0438"]
    assert sheet.freeze_panes == "A2"
    assert sheet["A1"].value == "\u0418\u043d\u0432"
    assert sheet["B1"].value == "\u0421\u0435\u0440\u0438\u0439\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440"
    assert sheet["A1"].font.bold is True
    assert sheet["A2"].value == "101"
    assert sheet["B2"].value == "\u2014"
    assert sheet["C2"].value == "Dell Latitude"
    uploads[0].file.close()


def test_ai_files_report_tool_accepts_top_level_rows_and_fallback_fields():
    from openpyxl import load_workbook

    from backend.ai_chat.artifact_generator import build_generated_uploads
    from backend.ai_chat.tools.files import FilesReportArgs, FilesReportTool

    result = FilesReportTool().execute(
        context=_make_context(),
        args=FilesReportArgs.model_validate({
            "file_name": "",
            "title": "",
            "rows": [["\u0418\u043d\u0432", "\u041c\u043e\u0434\u0435\u043b\u044c"], ["101", "Dell"]],
            "metadata": "ignored",
        }),
    ).to_payload()

    assert result["ok"] is True
    spec = result["data"]["files"][0]
    assert spec["file_name"] == "report.xlsx"
    assert len(spec["sheets"]) == 1
    assert spec["sheets"][0]["title"] == "Data"

    uploads = build_generated_uploads([spec])
    workbook = load_workbook(uploads[0].file, read_only=True)
    assert workbook.sheetnames[-1] == "Data"
    assert workbook["Data"].max_row == 2
    assert workbook["Data"].max_column == 3
    uploads[0].file.close()


def test_ai_files_report_tool_skips_empty_xlsx_tables_and_allows_large_tables():
    from openpyxl import load_workbook

    from backend.ai_chat.artifact_generator import build_generated_uploads
    from backend.ai_chat.tools.files import FilesReportArgs, FilesReportTool

    result = FilesReportTool().execute(
        context=_make_context(),
        args=FilesReportArgs.model_validate({
            "format": "xlsx",
            "file_name": "large-report",
            "title": "Large report",
            "summary": "Large export.",
            "tables": [
                {"title": "Empty tail", "rows": [{}, [], {"inv_no": "", "model_name": ""}]},
                {
                    "title": "Equipment",
                    "rows": [
                        {"inv_no": str(index), "serial_no": f"SN-{index}", "model_name": "Model"}
                        for index in range(1, 5002)
                    ],
                },
            ],
        }),
    ).to_payload()

    assert result["ok"] is True
    spec = result["data"]["files"][0]
    assert [sheet["title"] for sheet in spec["sheets"]] == ["\u0421\u0432\u043e\u0434\u043a\u0430", "Equipment"]
    assert len(spec["sheets"][1]["rows"]) == 5002

    uploads = build_generated_uploads([spec])
    workbook = load_workbook(uploads[0].file, read_only=True)
    assert workbook.sheetnames == ["\u0421\u0432\u043e\u0434\u043a\u0430", "Equipment"]
    assert workbook["Equipment"].max_row == 5002
    uploads[0].file.close()


def test_ai_files_report_tool_honors_explicit_table_columns():
    from backend.ai_chat.tools.files import FilesReportArgs, FilesReportTool

    result = FilesReportTool().execute(
        context=_make_context(),
        args=FilesReportArgs.model_validate({
            "format": "csv",
            "file_name": "ordered",
            "title": "Ordered",
            "tables": [
                {
                    "title": "Data",
                    "columns": [
                        {"key": "model", "label": "\u041c\u043e\u0434\u0435\u043b\u044c"},
                        {"key": "inv", "label": "\u0418\u043d\u0432"},
                    ],
                    "rows": [{"inv": "101", "model": "Dell"}],
                }
            ],
        }),
    ).to_payload()

    assert result["ok"] is True
    spec = result["data"]["files"][0]
    assert spec["file_name"] == "ordered.csv"
    assert spec["rows"] == [
        ["\u041c\u043e\u0434\u0435\u043b\u044c", "\u0418\u043d\u0432", "\u0421\u0435\u0440\u0438\u0439\u043d\u044b\u0439 \u043d\u043e\u043c\u0435\u0440"],
        ["Dell", "101", "\u2014"],
    ]


def test_ai_files_report_tool_builds_docx_tables_from_report_structure():
    from docx import Document

    from backend.ai_chat.artifact_generator import build_generated_uploads
    from backend.ai_chat.tools.files import FilesReportArgs, FilesReportTool

    result = FilesReportTool().execute(
        context=_make_context(),
        args=FilesReportArgs.model_validate({
            "format": "docx",
            "file_name": "table-report",
            "title": "\u041e\u0442\u0447\u0451\u0442",
            "summary": "\u0418\u0442\u043e\u0433.",
            "tables": [
                {
                    "title": "\u0422\u0435\u0445\u043d\u0438\u043a\u0430",
                    "rows": [["\u0418\u043d\u0432", "\u041c\u043e\u0434\u0435\u043b\u044c"], ["101", "Dell"]],
                }
            ],
        }),
    ).to_payload()

    assert result["ok"] is True
    uploads = build_generated_uploads([result["data"]["files"][0]])
    document = Document(uploads[0].file)
    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "\u0418\u043d\u0432"
    assert document.tables[0].cell(1, 1).text == "\u2014"
    assert document.tables[0].cell(1, 2).text == "Dell"
    uploads[0].file.close()


def test_ai_files_report_tool_builds_pdf_from_report_structure():
    from backend.ai_chat.artifact_generator import build_generated_uploads
    from backend.ai_chat.tools.files import FilesReportArgs, FilesReportTool

    result = FilesReportTool().execute(
        context=_make_context(),
        args=FilesReportArgs.model_validate({
            "format": "pdf",
            "file_name": "table-report",
            "title": "\u041e\u0442\u0447\u0451\u0442",
            "summary": "\u0420\u0443\u0441\u0441\u043a\u0438\u0439 \u0442\u0435\u043a\u0441\u0442.",
            "tables": [{"title": "Data", "rows": [["A", "B"], ["1", "2"]]}],
        }),
    ).to_payload()

    assert result["ok"] is True
    uploads = build_generated_uploads([result["data"]["files"][0]])
    assert uploads[0].filename == "table-report.pdf"
    assert uploads[0].file.getbuffer().nbytes > 1000
    uploads[0].file.close()


def test_ai_files_create_tool_returns_structured_diagnostic_on_generator_error():
    from backend.ai_chat.tools.files import FilesCreateArgs, FilesCreateTool

    result = FilesCreateTool().execute(
        context=_make_context(),
        args=FilesCreateArgs.model_validate({
            "files": [{"format": "csv", "file_name": "bad", "rows": ["not", "a", "table"]}]
        }),
    ).to_payload()

    assert result["ok"] is False
    assert result["data"]["diagnostic"]["error_code"] == "invalid_rows_shape"
    assert result["data"]["diagnostic"]["field_path"] == "rows"


def test_ai_files_report_tool_respects_generated_files_master_switch():
    from backend.ai_chat.tools.files import FilesReportArgs, FilesReportTool

    result = FilesReportTool().execute(
        context=_make_context(allow_generated_artifacts=False),
        args=FilesReportArgs.model_validate({
            "format": "pdf",
            "file_name": "blocked.pdf",
            "title": "Blocked",
            "summary": "Blocked",
        }),
    ).to_payload()

    assert result["ok"] is False
    assert "disabled" in result["error"]
